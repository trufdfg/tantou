# -*- coding: utf-8 -*-
from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from PIL import Image, ImageDraw, ImageFont
import math

OUT = Path('artifacts')
OUT.mkdir(exist_ok=True)
IMG = OUT / 'figures'
IMG.mkdir(exist_ok=True)
DOCX = OUT / '基于亚太赫兹侧壁回波的电缆充油终端分层识别与界面定位_硕士开题报告_插图完整版.docx'

FONT_CJK = '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc'
FONT_CJK_B = '/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc'
if not Path(FONT_CJK_B).exists(): FONT_CJK_B = FONT_CJK

def fnt(size, bold=False):
    return ImageFont.truetype(FONT_CJK_B if bold else FONT_CJK, size)

BLUE=(21,77,134); BLUE2=(45,111,170); TEAL=(26,137,139); GREEN=(55,142,87); ORANGE=(230,126,34)
LIGHT_BLUE=(232,242,252); LIGHT_GREEN=(232,247,238); LIGHT_ORANGE=(255,244,229); GRAY=(90,100,110); DARK=(35,45,55)

def rounded(draw, xy, r=18, fill='white', outline=BLUE, width=3):
    draw.rounded_rectangle(xy, radius=r, fill=fill, outline=outline, width=width)

def text_center(draw, xy, text, font, fill=DARK, spacing=8):
    box=draw.multiline_textbbox((0,0), text, font=font, spacing=spacing, align='center')
    w=box[2]-box[0]; h=box[3]-box[1]
    draw.multiline_text(((xy[0]+xy[2]-w)/2,(xy[1]+xy[3]-h)/2), text, font=font, fill=fill, spacing=spacing, align='center')

def arrow(draw, a, b, color=BLUE, width=7):
    draw.line([a,b], fill=color, width=width)
    ang=math.atan2(b[1]-a[1], b[0]-a[0]); L=18
    p1=(b[0]-L*math.cos(ang-0.55), b[1]-L*math.sin(ang-0.55))
    p2=(b[0]-L*math.cos(ang+0.55), b[1]-L*math.sin(ang+0.55))
    draw.polygon([b,p1,p2], fill=color)

def save(img, name):
    p=IMG/name; img.save(p, dpi=(220,220)); return p

# Figure 1: technical route
img=Image.new('RGB',(1800,1050),'white'); d=ImageDraw.Draw(img)
text_center(d,(0,15,1800,100),'图1  课题总体技术路线',fnt(48,True),(15,35,70))
cols=[('1 研究目标与问题定义',['空气/硅油/水识别','双界面定位','厚度/距离/角度边界'],BLUE,LIGHT_BLUE),
      ('2 材料与传播机理',['公开参数与THz-TDS趋势','多层Fresnel/传输矩阵','目标工作频点校准'],TEAL,(231,248,248)),
      ('3 仿真与前期验证',['二维/三维COMSOL','幅值/频率/介电参数扫描','已完成空气/硅油机理验证'],GREEN,LIGHT_GREEN),
      ('4 实验平台与数据集',['中心频率约122 GHz探头','RS485：220点+CFAR','模拟终端轴向扫描'],ORANGE,LIGHT_ORANGE),
      ('5 识别与界面定位',['门控与比值特征','三介质概率分类','变点与概率交点定位'],BLUE2,LIGHT_BLUE),
      ('6 影响因素与应用',['外壳厚度/站距/角度','温度与表面凝露','STM32与真实终端适用性'],TEAL,(231,248,248))]
left=28; gap=18; cw=(1800-2*left-5*gap)//6
for i,(title,items,c,fc) in enumerate(cols):
    x=left+i*(cw+gap); rounded(d,(x,135,x+cw,835),22,'white',c,4)
    d.rounded_rectangle((x,135,x+cw,250),radius=22,fill=c)
    d.rectangle((x,220,x+cw,250),fill=c)
    text_center(d,(x+8,145,x+cw-8,245),title,fnt(27,True),'white')
    y=290
    for item in items:
        d.ellipse((x+22,y+8,x+36,y+22),fill=c)
        d.multiline_text((x+48,y),item,font=fnt(23),fill=DARK,spacing=6)
        y+=145
    if i in (2,3):
        tag='已完成基础' if i==2 else '已完成初步链路'
        d.rounded_rectangle((x+35,720,x+cw-35,785),radius=20,fill=(225,247,229),outline=GREEN,width=2)
        text_center(d,(x+35,720,x+cw-35,785),tag,fnt(23,True),GREEN)
    if i<5: arrow(d,(x+cw+2,480),(x+cw+gap-2,480),BLUE,6)
rounded(d,(90,885,1710,1015),28,(245,249,253),BLUE,4)
text_center(d,(110,900,1690,1000),'形成“材料参数—传播机理—仿真—实验—识别定位—适用边界—工程验证”闭环研究体系',fnt(33,True),BLUE)
fig1=save(img,'fig1_technical_route.png')

# Figure 2: principle
img=Image.new('RGB',(1800,1100),'white'); d=ImageDraw.Draw(img)
text_center(d,(0,10,1800,95),'图2  亚太赫兹侧壁检测原理与分层界面定位示意',fnt(45,True),(15,35,70))
# probe
rounded(d,(70,365,390,600),35,(235,240,246),BLUE,5)
text_center(d,(85,385,375,535),'亚太赫兹FMCW探头\n中心工作频率约122 GHz',fnt(30,True),BLUE)
for k,c in enumerate([BLUE,TEAL,ORANGE]):
    d.arc((330-k*5,410-k*25,520+k*40,565+k*25),-48,48,fill=c,width=6)
# termination
cx=1030; top=150; bottom=930; outer=450; inner=330
# outer shell
rounded(d,(cx-outer//2,top,cx+outer//2,bottom),70,(235,237,240),(80,85,90),6)
rounded(d,(cx-inner//2,top+55,cx+inner//2,bottom-55),45,(235,246,252),(100,105,110),4)
# layers
inner_x1=cx-inner//2+12; inner_x2=cx+inner//2-12; y1=top+70; y2=bottom-70
h=(y2-y1)
a_end=y1+int(h*0.31); o_end=y1+int(h*0.72)
d.rectangle((inner_x1,y1,inner_x2,a_end),fill=(226,242,252))
d.rectangle((inner_x1,a_end,inner_x2,o_end),fill=(252,229,139))
d.rectangle((inner_x1,o_end,inner_x2,y2),fill=(159,218,244))
text_center(d,(inner_x1,a_end-120,inner_x2,a_end-20),'空气层',fnt(35,True),BLUE)
text_center(d,(inner_x1,o_end-150,inner_x2,o_end-35),'硅油层',fnt(35,True),(145,95,0))
text_center(d,(inner_x1,y2-140,inner_x2,y2-35),'水层',fnt(35,True),BLUE)
# conductor
d.rounded_rectangle((cx-35,top-40,cx+35,bottom+45),radius=24,fill=(190,103,40),outline=(110,60,20),width=4)
# interfaces
d.line((inner_x1,a_end,inner_x2,a_end),fill=ORANGE,width=6)
d.line((inner_x1,o_end,inner_x2,o_end),fill=BLUE,width=6)
d.text((cx+250,a_end-25),'空气—硅油界面',font=fnt(28,True),fill=ORANGE)
d.text((cx+250,o_end-25),'硅油—水界面',font=fnt(28,True),fill=BLUE)
# echo paths
for yy,col,label in [(345,BLUE,'外表面回波'),(500,TEAL,'内壁回波'),(650,(115,70,170),'多次回波')]:
    arrow(d,(420,yy),(inner_x1+10,yy),col,6); arrow(d,(inner_x1+10,yy+18),(430,yy+18),col,5)
    d.text((470,yy-55),label,font=fnt(24,True),fill=col)
# scan arrow
arrow(d,(1510,830),(1510,230),BLUE,7); arrow(d,(1510,230),(1510,830),BLUE,7)
d.text((1550,475),'沿轴向\n连续扫描',font=fnt(30,True),fill=BLUE,spacing=8)
rounded(d,(80,945,1720,1055),25,(245,249,253),BLUE,3)
text_center(d,(90,955,1710,1045),'单点侧壁回波用于判断局部介质类别；类别概率沿高度方向的跃迁用于定位分层界面',fnt(31,True),BLUE)
fig2=save(img,'fig2_principle.png')

# Figure 3: simulation completed work
img=Image.new('RGB',(1800,1020),'white'); d=ImageDraw.Draw(img)
text_center(d,(0,10,1800,95),'图3  已完成的仿真工作与后续校准关系',fnt(45,True),(15,35,70))
blocks=[(70,150,555,850,'已完成：二维/三维机理仿真',GREEN,LIGHT_GREEN),
        (655,150,1140,850,'已完成：参数扫描与特征分析',TEAL,(231,248,248)),
        (1240,150,1730,850,'后续：目标频点闭环校准',ORANGE,LIGHT_ORANGE)]
for x1,y1,x2,y2,title,c,fc in blocks:
    rounded(d,(x1,y1,x2,y2),25,'white',c,4)
    d.rounded_rectangle((x1,y1,x2,y1+105),radius=25,fill=c); d.rectangle((x1,y1+75,x2,y1+105),fill=c)
    text_center(d,(x1+10,y1+8,x2-10,y1+98),title,fnt(27,True),'white')
# model schematic
x1,y1,x2,y2=110,300,510,650
d.rectangle((x1,y1,x2,y2),fill=(230,234,238),outline=GRAY,width=3)
d.rectangle((x1+70,y1+40,x2-70,y1+150),fill=(225,242,252))
d.rectangle((x1+70,y1+150,x2-70,y2-40),fill=(252,229,139))
arrow(d,(130,475),(470,475),BLUE,8)
d.text((145,330),'空气区域',font=fnt(27,True),fill=BLUE); d.text((145,520),'硅油区域',font=fnt(27,True),fill=(145,95,0))
d.text((105,700),'COMSOL局部传播域\n高斯调制脉冲、面积平均电场',font=fnt(25),fill=DARK,spacing=8)
# curves
for idx,col in enumerate([BLUE,TEAL,ORANGE]):
    pts=[]
    for xx in range(710,1090,4):
        t=(xx-710)/70
        amp=(1.0+idx*0.25)*math.exp(-0.22*t)*math.sin(2.5*t+idx*0.5)
        yy=430+int(85*amp)
        pts.append((xx,yy))
    d.line(pts,fill=col,width=4)
d.line((710,430,1090,430),fill=(180,180,180),width=2)
d.text((715,285),'扫描变量：',font=fnt(27,True),fill=TEAL)
d.multiline_text((715,335),'• 入射电场幅值\n• 中心频率\n• 硅油复介电参数\n• 老化/损耗等效参数',font=fnt(25),fill=DARK,spacing=12)
d.multiline_text((715,610),'已形成：时间门控、峰值、\n峰—峰值、能量积分和比值判据',font=fnt(25),fill=DARK,spacing=10)
# calibration
text_center(d,(1280,295,1690,440),'公开宽带参数\n+ THz-TDS频谱趋势',fnt(28,True),ORANGE)
arrow(d,(1485,455),(1485,550),ORANGE,7)
text_center(d,(1280,555,1690,690),'中心工作频率约122 GHz\n标准样件与实测回波反演',fnt(28,True),ORANGE)
arrow(d,(1485,700),(1485,785),ORANGE,7)
text_center(d,(1280,790,1690,835),'建立仿真—实测对应关系',fnt(25,True),ORANGE)
rounded(d,(140,900,1660,985),22,(245,249,253),BLUE,3)
text_center(d,(150,905,1650,980),'现有1–10 THz机理仿真用于证明物理可行性，不直接等同于122 GHz实物探头定量结论',fnt(28,True),BLUE)
fig3=save(img,'fig3_simulation_basis.png')

# Figure 4: experimental platform and completed work
img=Image.new('RGB',(1800,1080),'white'); d=ImageDraw.Draw(img)
text_center(d,(0,10,1800,95),'图4  已完成的模拟终端实物链路及后续平台升级',fnt(45,True),(15,35,70))
# left platform
rounded(d,(45,135,720,930),28,'white',BLUE,4)
d.rounded_rectangle((45,135,720,245),radius=28,fill=BLUE); d.rectangle((45,215,720,245),fill=BLUE)
text_center(d,(55,145,710,235),'已完成：模拟终端与有油/无油检测',fnt(29,True),'white')
# tube
rounded(d,(165,300,385,760),55,(245,246,235),(130,125,90),5)
d.rectangle((185,310,365,500),fill=(238,242,230)); d.rectangle((185,500,365,745),fill=(249,223,118))
d.text((230,380),'空气',font=fnt(30,True),fill=BLUE); d.text((215,610),'硅油',font=fnt(30,True),fill=(145,95,0))
# probe and board
rounded(d,(430,445,650,575),26,(225,232,240),TEAL,4); text_center(d,(440,455,640,565),'亚太赫兹探头',fnt(25,True),TEAL)
for k in range(3): d.arc((370-k*12,470-k*20,455+k*5,555+k*20),-55,55,fill=TEAL,width=5)
rounded(d,(430,635,650,760),15,(220,237,248),BLUE,3); text_center(d,(440,642,640,750),'STM32F103\n红灯：有油\n绿灯：无油',fnt(23,True),BLUE)
d.multiline_text((90,800),'模拟结构：玻璃纤维 + 硅橡胶管\n已打通探头—RS485—单片机判定链路',font=fnt(25),fill=DARK,spacing=10)
# middle data
rounded(d,(760,135,1240,930),28,'white',TEAL,4)
d.rounded_rectangle((760,135,1240,245),radius=28,fill=TEAL); d.rectangle((760,215,1240,245),fill=TEAL)
text_center(d,(770,145,1230,235),'已完成：数据采集与初步处理',fnt(29,True),'white')
d.multiline_text((820,295),'探头接口：RS485，115200 bps\n单帧输出：Pos、Dis、220点幅值、CFAR\n采集方式：多周向位置、分距离重复采集',font=fnt(24),fill=DARK,spacing=16)
# mini waveform
for series,col,offset in [(0,BLUE,0),(1,ORANGE,35)]:
    pts=[]
    for xx in range(815,1185,4):
        t=(xx-815)/22
        yy=650+offset-int(70*math.exp(-0.18*t)*math.sin(1.8*t+series))
        pts.append((xx,yy))
    d.line(pts,fill=col,width=4)
d.line((810,720,1190,720),fill=GRAY,width=2); d.line((810,530,810,720),fill=GRAY,width=2)
d.text((840,745),'已完成波形绘制、基础特征与阈值/模型初判',font=fnt(23),fill=TEAL)
# counts
rounded(d,(825,805,1175,885),18,(232,247,238),GREEN,3)
text_center(d,(830,810,1170,880),'无油 482 组；有油 461 组\n每组 220 点',fnt(25,True),GREEN)
# right upgrade
rounded(d,(1280,135,1755,930),28,'white',ORANGE,4)
d.rounded_rectangle((1280,135,1755,245),radius=28,fill=ORANGE); d.rectangle((1280,215,1755,245),fill=ORANGE)
text_center(d,(1290,145,1745,235),'后续：标准化平台与定量界面定位',fnt(28,True),'white')
items=['可观察液位真值的透明内胆','可更换同材质外壳厚度','轴向导轨与角度/站距夹具','空气—硅油—水三层结构','正反向扫描与独立盲测','跨日/跨安装数据集划分']
y=300
for i,it in enumerate(items,1):
    d.ellipse((1320,y+4,1350,y+34),fill=ORANGE); text_center(d,(1320,y+4,1350,y+34),str(i),fnt(18,True),'white')
    d.text((1370,y),it,font=fnt(24),fill=DARK); y+=90
rounded(d,(180,960,1620,1045),20,(245,249,253),BLUE,3)
text_center(d,(190,965,1610,1040),'前期实物结果证明有限场景下有油/无油具有可分性，但尚不能替代界面精度、鲁棒性和真实终端验证',fnt(27,True),BLUE)
fig4=save(img,'fig4_experiment_basis.png')

# Figure 5: algorithm
img=Image.new('RGB',(1800,1000),'white'); d=ImageDraw.Draw(img)
text_center(d,(0,10,1800,95),'图5  回波处理、三介质识别与界面定位流程',fnt(45,True),(15,35,70))
steps=[('1 原始输入','220点幅值\nCFAR序列\nPos/Dis与元数据',BLUE,LIGHT_BLUE),
       ('2 质量控制与预处理','帧完整性检查\n背景扣除\n外表面峰对齐\n门控归一化',TEAL,(231,248,248)),
       ('3 物理特征构建','峰值/峰宽\n门控能量比\n多次回波衰减\n波形统计量',GREEN,LIGHT_GREEN),
       ('4 介质识别','逻辑回归/LDA\nSVM/随机森林\n输出空气/硅油/水概率',ORANGE,LIGHT_ORANGE),
       ('5 界面定位','概率交点\n变点检测\n空气—硅油界面\n硅油—水界面',BLUE2,LIGHT_BLUE),
       ('6 评价与部署','F1/MAE/置信区间\n厚度/距离/角度边界\n低置信拒识\nSTM32一致性',TEAL,(231,248,248))]
left=25; gap=18; cw=(1800-50-5*gap)//6
for i,(title,body,c,fc) in enumerate(steps):
    x=left+i*(cw+gap); rounded(d,(x,150,x+cw,825),22,'white',c,4)
    d.rounded_rectangle((x,150,x+cw,270),radius=22,fill=c); d.rectangle((x,235,x+cw,270),fill=c)
    text_center(d,(x+8,160,x+cw-8,260),title,fnt(26,True),'white')
    text_center(d,(x+18,300,x+cw-18,610),body,fnt(24),DARK,spacing=14)
    # simple icon panel
    d.rounded_rectangle((x+45,650,x+cw-45,765),radius=18,fill=fc,outline=c,width=2)
    if i==0:
        pts=[(x+60+j*5,720-int(35*math.exp(-0.06*j)*math.sin(j/3))) for j in range(0,30)]
        d.line(pts,fill=c,width=4)
    elif i==1:
        d.polygon([(x+75,675),(x+cw-75,675),(x+cw-110,730),(x+110,730)],fill=c)
    elif i==2:
        for j,h in enumerate([35,65,45,85]): d.rectangle((x+70+j*38,745-h,x+95+j*38,745),fill=c)
    elif i==3:
        d.ellipse((x+70,690,x+115,735),fill=BLUE); d.polygon([(x+145,735),(x+170,690),(x+195,735)],fill=GREEN); d.rectangle((x+215,700,x+255,740),fill=ORANGE)
    elif i==4:
        d.line((x+60,730,x+cw-60,730),fill=c,width=3); d.line((x+100,680,x+100,750),fill=ORANGE,width=4); d.line((x+cw-100,680,x+cw-100,750),fill=BLUE,width=4)
    else:
        d.rounded_rectangle((x+90,675,x+cw-90,745),radius=12,outline=c,width=5); d.text((x+cw//2-32,691),'MCU',font=fnt(23,True),fill=c)
    if i<5: arrow(d,(x+cw+2,490),(x+cw+gap-2,490),BLUE,6)
rounded(d,(150,875,1650,965),22,(245,249,253),BLUE,3)
text_center(d,(160,880,1640,960),'已完成有油/无油初步判别链路；后续重点扩展三介质概率识别、双界面定位与适用边界评价',fnt(29,True),BLUE)
fig5=save(img,'fig5_algorithm.png')

# -------- document helpers --------
doc=Document()
sec=doc.sections[0]
sec.top_margin=Cm(2.5); sec.bottom_margin=Cm(2.3); sec.left_margin=Cm(2.7); sec.right_margin=Cm(2.5)

styles=doc.styles
styles['Normal'].font.name='Times New Roman'; styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体'); styles['Normal'].font.size=Pt(12)
for name,size,cn in [('Title',22,'黑体'),('Heading 1',16,'黑体'),('Heading 2',14,'黑体'),('Heading 3',12,'黑体')]:
    st=styles[name]; st.font.name='Times New Roman'; st._element.rPr.rFonts.set(qn('w:eastAsia'),cn); st.font.size=Pt(size); st.font.bold=True
styles['Heading 1'].paragraph_format.space_before=Pt(12); styles['Heading 1'].paragraph_format.space_after=Pt(8)
styles['Heading 2'].paragraph_format.space_before=Pt(8); styles['Heading 2'].paragraph_format.space_after=Pt(5)

# custom caption
if '图题' not in styles:
    st=styles.add_style('图题',WD_STYLE_TYPE.PARAGRAPH); st.font.name='Times New Roman'; st._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体'); st.font.size=Pt(10.5)

def set_cell_shading(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)

def set_cell_text(cell,text,bold=False,color=None,size=10.5):
    cell.text=''; p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); r.bold=bold; r.font.size=Pt(size); r.font.name='Times New Roman'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
    if color: r.font.color.rgb=RGBColor(*color)
    cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_body(text,bold_lead=None):
    p=doc.add_paragraph(); p.paragraph_format.first_line_indent=Cm(0.74); p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(4); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_lead and text.startswith(bold_lead):
        r=p.add_run(bold_lead); r.bold=True; r._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体'); p.add_run(text[len(bold_lead):])
    else: p.add_run(text)
    return p

def add_bullet(text):
    p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.line_spacing=1.35; p.paragraph_format.space_after=Pt(2); p.add_run(text); return p

def add_heading(text,level=1):
    p=doc.add_heading(text,level=level); p.paragraph_format.keep_with_next=True; return p

def add_figure(path,caption,width=15.8):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(5); p.paragraph_format.space_after=Pt(2)
    p.add_run().add_picture(str(path),width=Cm(width))
    c=doc.add_paragraph(style='图题'); c.alignment=WD_ALIGN_PARAGRAPH.CENTER; c.add_run(caption)

def add_table(headers, rows, widths=None, header_fill='1F4E78'):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
    for i,h in enumerate(headers): set_cell_text(t.rows[0].cells[i],h,True,(255,255,255),10); set_cell_shading(t.rows[0].cells[i],header_fill)
    for row in rows:
        cells=t.add_row().cells
        for i,val in enumerate(row): set_cell_text(cells[i],str(val),False,None,9.5)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Cm(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)
    return t

def page_break(): doc.add_page_break()

# header/footer
header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.CENTER
rr=header.add_run('基于亚太赫兹侧壁回波的电缆充油终端分层识别与界面定位研究'); rr.font.size=Pt(9); rr.font.color.rgb=RGBColor(100,100,100); rr._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
run=footer.add_run(); fldChar1=OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'),'begin'); instr=OxmlElement('w:instrText'); instr.set(qn('xml:space'),'preserve'); instr.text='PAGE'; fldChar2=OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'),'end'); run._r.append(fldChar1); run._r.append(instr); run._r.append(fldChar2)

# Cover
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(28)
r=p.add_run('硕士研究生学位论文开题报告'); r.bold=True; r.font.size=Pt(20); r._element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(24); p.paragraph_format.space_after=Pt(14)
r=p.add_run('基于亚太赫兹侧壁回波的电缆充油终端\n空气—硅油—水分层状态识别与界面定位研究'); r.bold=True; r.font.size=Pt(24); r._element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Sub-Terahertz Sidewall-Echo-Based Identification and Interface Localization\nof Air–Silicone Oil–Water Stratification in Oil-Filled Cable Terminations'); r.font.size=Pt(12); r.italic=True
for _ in range(3): doc.add_paragraph()
info=[('学　　院',''),('专　　业','电气工程'),('研究方向','高电压与绝缘技术 / 电气设备状态检测'),('研究生',''),('指导教师',''),('日　　期','2026 年 7 月')]
t=doc.add_table(rows=len(info),cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,(a,b) in enumerate(info):
    t.rows[i].cells[0].width=Cm(4); t.rows[i].cells[1].width=Cm(8); set_cell_text(t.rows[i].cells[0],a,False,None,12); set_cell_text(t.rows[i].cells[1],b,False,None,12)
    for c in t.rows[i].cells:
        tcPr=c._tc.get_or_add_tcPr(); borders=OxmlElement('w:tcBorders'); tcPr.append(borders)
for _ in range(2): doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('研究主线统一表述为“亚太赫兹”，具体实验设备在实验部分说明中心工作频率约 122 GHz。').italic=True
page_break()

# Abstract
add_heading('摘　要',1)
add_body('电缆充油终端是高压电缆系统中电场、热场和机械应力高度集中的关键附件。硅油泄漏、油位下降、密封失效及水分侵入会改变终端内部绝缘介质分布和电场边界，若不能及时发现，可能诱发局部放电、异常发热、沿面放电乃至绝缘击穿。现有红外、超声、射线及常规电气诊断方法各有适用条件，但在不拆解、不开孔、不接触内部介质的条件下，从复合绝缘侧壁直接识别空气、硅油和水，并定量定位分层界面，仍缺少系统研究。')
add_body('本课题以亚太赫兹侧壁回波为信息载体，研究电缆充油终端内部空气—硅油—水介质状态的可观测性、可识别性和分层界面的可定位性。实验中选用中心工作频率约 122 GHz 的线性调频连续波亚太赫兹探头，通过 RS485 接口获取 Pos、Dis、220 点回波幅值和 CFAR 阈值序列，沿终端轴向进行侧壁扫描，重点回答“利用现有回波输出能够测到什么、界面能够定位多准、在何种结构与安装条件下仍然准确”三个核心问题。')
add_body('前期已完成侧壁传播机理分析、COMSOL 二维与局部三维瞬态电磁仿真、入射电场幅值和中心频率等参数扫描，并建立时间门控、峰值、峰—峰值、能量积分及比值判据。实物方面，已在玻璃纤维与硅橡胶管构成的模拟终端上完成空气/硅油初步检测，打通探头—RS485—PC—STM32 数据链路；累计获得无油 482 组、有油 461 组回波，每组 220 点，并实现有油红灯、无油绿灯的离线初步判断。上述结果证明有限场景下具有初步可分性，但尚不能替代三介质识别、界面定位精度、影响因素边界及真实终端适用性验证。')
add_body('后续研究将以空气/硅油识别及空气—硅油界面定位为必做主线，以水介质识别、硅油—水界面定位和最小水层厚度为拓展内容；通过公开介电数据、THz-TDS 频谱趋势和目标工作频点实测反演约束材料参数，建立多层介质传播模型；搭建可观察真值、可更换壁厚、可调站距与角度的轴向扫描平台；采用外表面参考回波对齐、物理门控、回波比值、多次回波衰减和概率变点方法实现识别与定位；最终给出外壳厚度、探头站距、入射角及有限环境条件下的可靠检测域，并开展 STM32 离线实现与退役终端工程适用性分析。')
p=doc.add_paragraph(); p.paragraph_format.line_spacing=1.5; r=p.add_run('关键词：'); r.bold=True; p.add_run('电缆充油终端；亚太赫兹；侧壁回波；硅油液位；介质识别；界面定位；FMCW 雷达')
page_break()

# TOC manual
add_heading('目　录',1)
toc=[('1 选题背景与研究意义','4'),('2 国内外研究现状与发展趋势','8'),('3 研究计划','12'),('4 可行性论证与前期工作基础','25'),('5 研究进度安排','29'),('6 预期成果','30'),('7 拟定学位论文目录','31'),('参考文献','33')]
for title,num in toc:
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(6); p.paragraph_format.tab_stops.add_tab_stop(Cm(15.5)); p.add_run(title); p.add_run('\t'+num)
add_body('注：本报告参照博士研究生开题报告的“选题依据—研究现状—研究计划—可行性—进度与成果”逻辑组织，研究内容和考核指标按硕士学位论文可完成性进行约束。')
page_break()

# 1
add_heading('1 选题背景与研究意义',1)
add_heading('1.1 研究背景',2)
add_body('随着高压、超高压交联聚乙烯电缆线路持续增长，电缆附件已成为影响输电系统可靠性的薄弱环节之一。充油终端内部由贯通电缆、应力控制部件、硅橡胶绝缘、玻纤复合或瓷质外套、密封组件以及绝缘液体等构成，材料界面多、局部曲率大、安装离散性显著。制造缺陷、安装偏差、密封老化、材料受潮和运行热应力均可能演化为局部放电、异常发热和绝缘击穿，因此需要在故障形成早期获取内部状态信息。')
add_body('硅油用于填充终端内部空间、改善电场分布并辅助散热。密封圈老化、法兰松动、外壳裂纹或检修操作不当可能造成硅油缓慢泄漏，在上部形成空气区；外界水分还可能经密封薄弱处进入并在重力作用下形成高含水区、局部水滴或底部积水。空气区、硅油区和水区具有不同的介电强度、热传导与电磁损耗特征，三者分布变化不仅是液位问题，也是终端绝缘介质状态变化的重要表征。')
add_heading('1.2 工程检测需求与技术难点',2)
add_body('现场检测希望尽量满足不停电或少停电、不拆解、不在外壳开孔、不污染硅油以及可重复巡检等要求。传统顶部雷达需要对液面形成直接视线，而电缆终端中心导体贯通、顶部金属结构复杂，难以布置常规向下测距天线。本课题采用侧壁扫描：探头在某一高度从外侧照射终端，通过局部回波判断正对位置属于空气区、硅油区或水区；随后沿终端轴向移动，由介质类别沿高度方向的跃迁确定分层界面。')
add_bullet('强外表面反射可能掩盖穿透复合外壳后的内壁与介质回波。')
add_bullet('玻纤复合外壳和硅橡胶层厚度与亚太赫兹波长同量级，层内相位与多次反射可能发生增强或相消。')
add_bullet('探头站距与入射角的轻微变化可能引起比介质差异更大的幅值和峰位变化。')
add_bullet('液位坐标来自轴向扫描的类别概率过渡，而不是直接读取 FMCW 径向距离，其误差受扫描步长、波束足迹、弯月面与模型不确定度共同限制。')
add_heading('1.3 检测方法比较、亚太赫兹频段及实验载体选择',2)
add_body('红外热成像适合发现运行中的异常发热，但结果依赖负荷、电流、风速、日照和材料热传导；X 射线能够直观显示内部结构，但设备成本、辐射防护和现场布置限制了常态化应用；超声具有成熟的界面反射理论，但耦合层、接触压力、曲面和温度会影响重复性。亚太赫兹侧壁回波无需耦合剂、无电离辐射，可通过介电边界反射、吸收和多次回波直接感知局部介质状态，更适合作为现有手段的补充。')
add_body('亚太赫兹波位于微波与光子学太赫兹之间，兼具电子学系统的紧凑性与对介电参数差异的敏感性。实验选用中心工作频率约 122 GHz 的 FMCW 亚太赫兹探头，自由空间波长约 2.46 mm。该具体频率仅代表本课题的实验载体；研究主线与章节标题统一使用“亚太赫兹”，结论限定在该探头带宽、波束参数和输出接口条件下。FMCW 的名义径向距离分辨率主要由有效扫频带宽决定，不能把载频本身等同于液位定位精度。')
add_table(['方法','主要信息','优势','主要局限','本课题定位'],[
['红外热成像','表面温度场','非接触、巡检成熟','受负荷、风速、日照和热传导影响','文献比较'],
['X射线/数字射线','内部密度与结构投影','结构直观','电离辐射、成本与布置限制','高可信参照'],
['接触式超声/导波','声阻抗与回波能量','液—气边界敏感','耦合、曲面与温度影响','重点比较对象'],
['顶部雷达液位计','液面飞行时间','工业成熟','中心电缆贯通且顶部复杂','不采用'],
['亚太赫兹侧壁回波','介电反射、吸收与多次回波','无耦合剂、可侧壁扫描','受厚度、距离、角度影响','核心路线']], [2.5,3,3.1,4.3,2.8])
add_heading('1.4 研究意义',2)
add_body('理论意义。建立亚太赫兹波在“空气间隙—玻纤复合外壳—硅橡胶—空气/硅油/水”局部曲面多层结构中的传播模型，揭示外表面、层间、内壁与后续多次回波中可观测介质信息的形成机制。')
add_body('方法意义。将液位检测拆解为“单点局部介质识别”和“轴向类别跃迁定位”两个问题，构建参考回波归一化、物理门控、概率分类和变点定位的联合方法，并以不确定度和适用边界替代只报告单一工况准确率。')
add_body('工程意义。形成适用于不透明非金属侧壁的非接触、无电离辐射检测方案，验证低维物理特征与轻量模型在 RS485—STM32 离线链路上的实现能力。')

# 2
add_heading('2 国内外研究现状与发展趋势',1)
add_heading('2.1 电缆终端状态监测与缺陷检测研究',2)
add_body('电缆终端状态检测已形成局部放电、红外、超声、射线、电场测量、微波/太赫兹和多源融合等技术体系。现有方法普遍面临检测对象结构复杂、环境干扰强、定量评价不足及现场可重复性不高等问题。微波反射研究已经证明，多层介质中的气隙、界面错位和材料厚度变化会改变反射曲线，为利用电磁回波检测终端内部介电不连续提供了方法基础。')
add_heading('2.2 终端与密闭容器液位检测研究',2)
add_body('超声是从容器外检测液位的成熟路线。已有瓷套式终端研究发现，油区与空气区首界面回波差异可能较小，而多次回波衰减具有更稳定的辨识能力；侧壁声阻抗、导波泄漏能量及非接触 EMAT 方法也证明沿壁移动传感器可定位液位，但同时暴露出耦合、壁厚共振、频率选择与重复安装的约束。')
add_body('毫米波和微波研究表明，频率调制信号能够穿过不透明容器测量液位，利用容器前后界面的双反射比值能够削弱位置和容器影响，频域反射与干涉还能识别油—水多界面。这些研究多面向规则容器或能获得复数基带信号的系统，尚未解决电缆终端复合厚壁、中心导体散射和仅有 220 点幅值/CFAR 输出条件下的识别与定位问题。')
add_heading('2.3 太赫兹/亚太赫兹无损检测与多层介质测量',2)
add_body('太赫兹技术已用于 XLPE 电缆内部缺陷、复合绝缘子气隙与夹杂、油纸绝缘脱粘、GFRP 层厚和含水缺陷检测。相关研究说明玻纤复合材料具有一定可穿透性，多层结构中的 Fabry–Perot 干涉、脉冲重叠与材料吸收既可能造成识别困难，也能够提供层厚和损耗信息。大型 THz-TDS 适合获取材料频谱和相位，紧凑电子式 FMCW 亚太赫兹系统更适合现场检测，两者可形成“实验室频谱趋势—目标频点实测校准”的分层体系。')
add_heading('2.4 材料参数、水分与老化影响研究',2)
add_body('空气、硅油和水的可识别性来源于复介电参数差异。THz-TDS 能够获得复折射率、吸收系数与复介电常数，但样品厚度、平行度、时窗和 Fabry–Perot 效应会影响结果；实验室系统在 0.12 THz 附近的动态范围也可能不足。因此，本课题不将较高频段 TDS 参数机械外推至 122 GHz，而采用公开宽带模型、TDS 频谱趋势和标准样件实测反演进行闭环约束。')
add_heading('2.5 FMCW 回波处理与界面定位研究',2)
add_body('工业 FMCW 液位计通常通过拍频峰值估计目标距离，高精度算法多依赖原始复数基带、扫频相位或天线 S 参数。当前探头只输出处理后的 220 点幅值、CFAR 阈值及 Pos、Dis 信息，因此应将探头视为具有固定内部处理链的测量设备，通过标准反射体和位置扫描标定点序号、等效距离、幅值、重复性与波束足迹，优先使用幅值、峰形、门控能量和多次回波衰减。')
add_body('界面定位本质上是沿轴向的一维变点问题。界面附近由于波束同时覆盖两种介质，分类概率通常呈连续过渡，而非理想阶跃。定位误差由轴向步长、波束足迹、机械回差、界面弯月面和分类不确定度共同构成，需通过概率交点、单调拟合、正反向扫描和独立真值评价。')
add_heading('2.6 研究评述与本课题切入点',2)
for s in ['现有终端油位研究以超声、红外和射线为主，针对复合非金属侧壁的亚太赫兹三介质分层研究不足。','太赫兹电力绝缘研究多关注固体微缺陷、层厚和含水量，较少研究“复合外壳—硅橡胶—液体”曲面结构中的局部介质识别。','常规雷达液位研究以顶部测距为主，而电缆终端中心导体贯通，需要侧壁扫描和类别跃迁定位的新范式。','多数研究在单一工况下报告准确率，缺少外壳厚度、站距、角度变化后的可靠检测域、失效判据和低置信拒识。','从 PC 算法到 STM32 离线实现再到真实终端适用性分析的完整链条尚不充分。']:
    add_bullet(s)
add_body('据此，本课题以“现有回波输出的可观测性”为起点，以空气/硅油识别和空气—硅油界面定位为必做核心，以水介质与硅油—水界面为拓展，以厚度、距离、角度适用边界为主要评价，以前期仿真和实物链路为基础开展递进研究。')

# 3
add_heading('3 研究计划',1)
add_body('研究计划参照博士开题报告的组织逻辑，将研究目标、研究内容、理论方法、实验方案、关键问题与创新点纳入同一闭环；同时根据硕士课题的设备条件和已完成工作，对任务设置优先级，避免材料表征、全尺寸仿真、复杂深度模型和真实高压试验多线并行。')
add_heading('3.1 研究目标、核心问题与考核指标',2)
add_heading('3.1.1 总体研究目标',3)
add_body('面向电缆充油终端不透明非金属侧壁和内部贯通电缆结构，建立亚太赫兹侧壁回波检测方法，回答现有探头能够稳定观察哪些信息、空气—硅油—水能否可靠区分、双界面能够定位到何种精度，以及在何种外壳与安装条件下结果仍可信。')
add_heading('3.1.2 核心科学与技术问题',3)
add_table(['编号','核心问题','具体内涵'],[
['Q1','可观测性','强外表面回波与多层干涉背景下，220点幅值/CFAR中哪些门控区包含内部介质差异？'],
['Q2','可识别性','如何构造对系统增益、站距和曲率不敏感且适合单片机的低维物理特征？'],
['Q3','可定位性','如何把沿轴向离散测点的介质概率转换为界面位置并分离各误差来源？'],
['Q4','鲁棒性与边界','厚度、距离、角度和环境条件如何影响信号，何时应判为不可信？'],
['Q5','模型迁移','模拟终端模型如何迁移到退役真实终端，并控制校准样本与测试样本隔离？']], [1.5,3.1,10.5])
add_heading('3.1.3 预期技术指标',3)
add_table(['评价对象','建议目标','评价方式'],[
['空气/硅油识别','基准宏平均F1≥0.90；跨日/跨安装F1≥0.85','按容器、日期、安装批次分组留出，报告95%置信区间'],
['三介质识别','模拟平台宏平均F1≥0.85','水类别单独报告召回率与误报率'],
['空气—硅油界面','模拟终端MAE≤10 mm或不超过一个有效扫描步长','独立盲测，报告偏差、RMSE、最大误差'],
['硅油—水界面','模拟平台MAE≤15 mm','按水层厚度和界面位置分组'],
['适用边界','给出厚度—距离—角度可靠域','以F1、界面误差与质量分数共同定义'],
['嵌入式','PC与MCU同输入判定一致','报告运算时间、RAM/Flash占用']], [4.0,5.0,8.1])
add_heading('3.2 主要研究内容与任务分级',2)
add_table(['层级','研究任务','完成要求','研究边界'],[
['必做主线','空气/硅油局部识别；空气—硅油界面定位','跨日、跨容器、跨重新安装；报告F1、MAE和置信区间','对应漏油和低油位，构成最小完整闭环'],
['重点研究','外壳厚度、探头站距和入射角影响','建立主效应、交互趋势、补偿与可靠检测域','主要学术增量'],
['拓展内容','水介质识别、硅油—水界面、最小水层厚度','优先在透明模拟终端完成','真实终端积水不作强制要求'],
['辅助验证','THz-TDS趋势、有限三维仿真、STM32实现','参数约束、机理解释和工程实现','不开展全尺寸全参数三维扫频']], [2.3,4.4,5.1,5.3])
add_heading('3.2.1 材料频谱与目标工作频点参数校准',3)
add_body('采用“公开参数—THz-TDS 频谱趋势—目标工作频点实测反演”三级方式。对玻纤复合材料、硅橡胶和硅油测量频谱趋势与参数范围，对水采用宽带介电模型；在中心工作频率约 122 GHz 条件下，通过标准厚度样件、金属反射板和空场/材料回波校准等效参数。THz-TDS 只提供趋势与范围，不直接把 0.3–1 THz 或更高频率参数等同于目标频点。')
add_heading('3.2.2 多层介质传播与回波机理',3)
add_body('将局部结构简化为空气间隙—玻纤复合外壳—硅橡胶—内部介质，先用 Fresnel 系数与传输矩阵分析外表面、层间、内壁和多次回波，再用二维局部曲面模型修正斜入射、曲率与有限波束。重点识别现有探头能够观察的主导路径，而不是追求全尺寸终端的高成本全波重建。')
add_heading('3.2.3 三介质回波识别',3)
add_body('在同一外壳、测点与安装条件下分别构造空气、硅油和水单介质状态，再构建空气—硅油及空气—硅油—水分层结构。原始数据包括 220 点幅值、CFAR、Pos、Dis、液位真值、外壳厚度、站距、角度、温度和安装批次。预处理采用帧完整性检查、背景扣除、外表面峰对齐、物理门控、幅值归一化与距离补偿；特征优先选择外/内壁峰值比、门控能量比、多次回波衰减、峰宽、相关系数和 CFAR 超越面积。')
add_heading('3.2.4 分层界面定位与不确定度',3)
add_body('沿轴向以固定步长采集各位置的三类后验概率。空气—硅油界面取空气与硅油概率交点或对数似然比过零位置，硅油—水界面同理；界面区采用重叠扫描、逻辑曲线或单调样条拟合，并用物理顺序约束保证序列满足“空气→硅油→水”。定位真值由透明内胆标尺与独立位移测量给出，报告扫描机构、重复安装、概率拟合、波束足迹与弯月面造成的不确定度。')
add_heading('3.2.5 影响因素、适用边界与工程验证',3)
add_body('主要因素为外壳等效厚度、探头站距和入射角；温度与表面水膜/凝露作为工程验证因素。采用“单因素机理确认+有限多因素组合”策略，每个组合进行独立重新安装与完整轴向扫描。以识别 F1、界面 MAE、内壁门 SNR 和低置信样本比例共同定义可靠检测域；真实退役终端重点验证正常油位与低油位，不在带电或不可拆解设备中人为制造积水故障。')
add_figure(fig1,'图 1　课题总体技术路线',15.8)
add_heading('3.3 理论基础与研究方法',2)
add_heading('3.3.1 多层介质反射与传输模型',3)
add_body('对于非磁性介质，复相对介电常数可写为 εr*=ε′−jε″，复折射率为 n*=n−jκ，吸收系数 α=4πκ/λ0。法向入射时两介质界面的电场反射系数可近似写为 Γij=(ni*−nj*)/(ni*+nj*)。多层接收信号可表示为各传播路径的叠加 H(f)=ΣAp(f)exp(−j2πfτp)，其中 τp 由各层折射率和传播厚度共同决定。外壳厚度与波长同量级时，层内相位和 Fabry–Perot 干涉不可忽略，某些厚度窗口可能增强或抑制内部介质差异。')
add_heading('3.3.2 FMCW 输出标定与物理门控',3)
add_body('线性调频信号的拍频对应等效光程而非单纯几何距离。由于当前探头只输出处理后的幅值与 CFAR，应使用金属平板、吸波材料、标准介质和精密平移台建立“点序号—等效距离—幅值”的经验标定，测量开机预热漂移、短时重复性、重新安装再现性和轴向/周向波束足迹。信号划分为外表面参考门、内壁/介质响应门和后续多次回波门，以比值和能量分配降低系统增益与站距影响。')
add_heading('3.3.3 分类、域外拒识与模型解释',3)
add_body('分类模型按“可解释线性基线—非线性模型—深度模型对照”逐级比较，优先使用 LDA、逻辑回归、SVM 和随机森林；只有独立扫描序列数量充分时才引入 1D-CNN。域外拒识由输入质量门控、最大后验概率阈值和特征空间距离组成，对严重偏角、凝露、未知外壳或异常波形输出“低置信/需复测”，而不强制分类。')
add_figure(fig2,'图 2　亚太赫兹侧壁检测原理与分层界面定位示意',15.8)
add_heading('3.4 研究方案与实验设计',2)
add_heading('3.4.1 已完成仿真工作的继承与修正',3)
add_body('前期已基于典型充油终端结构建立二维和局部三维瞬态电磁模型，使用高斯调制脉冲研究侧壁传播，对空气区域与硅油区域的面积平均电场进行比较，并完成入射电场幅值、中心频率和硅油介电参数/老化等效参数扫描。现有结果已证明不同介质状态会导致可观测回波差异，时间门控、峰值、峰—峰值、能量积分和上下测点比值能够作为候选判据。')
add_body('但现有模型主要用于 1–10 THz 范围的机理验证，几何、激励形式和材料参数尚未与实际 122 GHz FMCW 探头完全对应。后续不将原结果直接作为目标频点定量结论，而是保留其多层传播机理与信号处理框架，补充一维传输矩阵和目标频点等效模型，并以标准样件和实测回波校准。')
add_figure(fig3,'图 3　已完成的仿真工作与目标工作频点校准关系',15.8)
add_heading('3.4.2 探头、数据采集与模拟终端',3)
add_body('实验采用中心工作频率约 122 GHz 的 FMCW 亚太赫兹探头，RS485 通信波特率 115200 bps，通过 !TEST:2 等指令输出 Pos、Dis、220 点幅值及对应 CFAR。PC 端保存完整帧和元数据，STM32F103 用于后期离线一致性验证。规范平台由透明内胆、可更换同材质等效外壳、硅橡胶层、中心贯通电缆等效件、轴向导轨、站距调节和俯仰/偏航角度夹具组成。')
add_heading('3.4.3 分阶段实验设计',3)
add_table(['阶段','主要任务','关键控制','阶段输出'],[
['A 测量链标定','金属板平移、角度扫描、预热与重复安装','固定反射板面积、法向基准、波束足迹','点序号—距离关系、重复性与质量门限'],
['B 单介质基准','空气、硅油、水重复采集','跨日、跨测点、重新安装、温度稳定','可观测门控区和候选特征'],
['C 分层定位','空气—硅油与三层轴向扫描','独立液位真值、正反向扫描、界面区小步长','概率序列、双界面位置、不确定度'],
['D 影响因素','厚度、站距、角度','先单因素后有限组合，每组合独立安装','主效应、交互作用和可靠域'],
['E 环境验证','温度、表面水膜/凝露','同步记录探头预热与表面状态','温度补偿、异常质量门控'],
['F 实际终端','退役终端正常/低油位盲测','测试人员与真值人员分离，测试集冻结','域偏移与工程适用性']], [2.3,4.8,5.3,5.3])
add_heading('3.4.4 数据集构建与防止数据泄漏',3)
add_body('同一 DAT 文件中的相邻帧高度相关，不能视为完全独立样本。实验统计单位至少为稳定测点窗口，更严格时为完整轴向扫描。训练、验证和测试按容器、日期、测点组、重新安装批次或真实终端分组；测试集在模型和阈值确定前冻结。最终同时报告帧级、测点级和整次扫描级性能。')
add_figure(fig4,'图 4　已完成的模拟终端实物链路及后续平台升级',15.8)
add_heading('3.4.5 回波处理、识别与界面定位算法',3)
add_body('算法输入与现有探头接口保持一致。首先检查 220 点帧完整性、异常行、饱和和重复帧；随后完成背景扣除、外表面峰对齐、物理门控和归一化；构建峰值比、能量比、多次回波衰减、峰形和统计特征；输出三类后验概率；最后利用概率交点、变点检测和顺序约束定位双界面。分类评价包括混淆矩阵、平衡准确率、宏平均 F1、各类别召回率与概率校准，定位评价包括偏差、MAE、RMSE、最大误差、95%误差范围和正反向扫描差异。')
add_figure(fig5,'图 5　回波处理、三介质识别与界面定位流程',15.8)
add_heading('3.5 拟解决的关键问题及应对措施',2)
add_table(['关键问题','主要风险','拟采取措施'],[
['外表面强回波掩盖内部信息','空气/硅油差异不可见或被固定杂波主导','标准反射板标定、外表面峰对齐、物理门控、背景差分和参考峰比值'],
['厚度引起干涉增强/相消','同一类别在不同厚度下特征方向反转','传输矩阵扫描敏感厚度，厚度作为模型输入，给出不可测窗口'],
['站距与角度变化大于介质差异','跨安装误判与概率失真','刚性夹具、距离补偿、角度增强训练、质量门控和低置信拒识'],
['界面附近波束混合','概率过渡宽、定位依赖步长','实测波束足迹、重叠扫描、概率拟合、正反向扫描和置信区间'],
['水层形态不理想','水滴、挂水或乳化导致标签不唯一','先研究静态平整水层，再设置挂水/水滴异常组并限定结论'],
['现有仿真与目标频点不匹配','机理结果被误作122 GHz定量结论','TDS与高频仿真只给趋势，使用目标频点标准样件与实测反演校准'],
['重复帧数据泄漏','测试准确率虚高','按文件/测点/日期/容器/终端分组，冻结盲测集']], [3.6,5.2,8.6])
add_heading('3.6 预期创新点',2)
add_body('（1）提出适用于中心电缆贯通结构的亚太赫兹侧壁分层检测范式。区别于常规顶部雷达直接测距，先利用局部径向回波识别空气、硅油和水，再由轴向类别概率跃迁定位界面。')
add_body('（2）建立与现有幅值/CFAR 输出匹配的“外表面参考—内壁响应—多次回波衰减”联合表征方法，将多次反射由单纯干扰项转化为介质吸收信息。')
add_body('（3）建立带物理顺序约束和不确定度评价的双界面定位方法，区分扫描步长、波束足迹、真值与分类模型对定位误差的贡献。')
add_body('（4）从单一条件识别扩展到厚度—距离—角度适用边界，形成可靠检测域、不可测窗口与低置信拒识规则。')
add_body('创新点最终以实验结果和持续文献检索为准，不使用“首次”“完全解决”等绝对表述。')

# 4
add_heading('4 可行性论证与前期工作基础',1)
add_heading('4.1 理论可行性',2)
add_body('空气、硅油和水具有不同的复介电常数、波阻抗与损耗，玻纤复合材料和硅橡胶属于可被亚太赫兹波一定程度穿透的绝缘材料。Fresnel 反射、传输矩阵、FMCW 距离像、门控能量与变点检测均有成熟理论基础；毫米波穿容器液位、太赫兹容器内液体表征与 GFRP 多层检测为本课题提供了直接方法依据。')
add_heading('4.2 技术与设备可行性',2)
add_body('现有条件包括亚太赫兹 FMCW 探头、RS485 数据接口、220 点幅值/CFAR 采集程序、STM32F103 开发板、TFT/LED 状态显示、学院 THz-TDS 设备以及 COMSOL 仿真基础。上述条件已覆盖“探头—采集—PC 处理—嵌入式判断”的基本链路，后续主要新增标准化轴向平台、可更换壁厚模拟终端、角度调节机构和独立真值系统，设备投入与加工难度可控。')
add_heading('4.3 已完成的仿真工作',2)
add_bullet('完成侧壁入射多层介质传播机理分析，明确方法本质是局部介质状态识别，而不是顶部飞行时间测距。')
add_bullet('完成 COMSOL 二维和局部三维瞬态电磁模型，采用高斯调制脉冲并提取接收面面积平均电场。')
add_bullet('完成空气区域与硅油区域对比，以及入射电场幅值、中心频率和硅油介电参数/老化等效参数扫描。')
add_bullet('形成时间门控、最大绝对幅值、峰—峰值、能量积分、上下测点幅值比和能量比等候选判据。')
add_body('前期仿真证明了空气/硅油状态产生回波差异的物理可行性，但所用频率和激励形式与实际 122 GHz FMCW 探头尚未完全一致。因此，新开题不要求推翻已有模型，而是将其定位为“机理与信号处理基础”，后续补充目标工作频点等效模型和实测闭环校准。')
add_heading('4.4 已完成的实物与数据工作',2)
add_bullet('在玻璃纤维与硅橡胶管组成的模拟终端上完成油区与空气区侧壁测量。')
add_bullet('完成 RS485 通信与 !TEST:2 指令测试，可稳定获得 Pos、Dis、220 点幅值和 CFAR 序列，通信波特率为 115200 bps。')
add_bullet('固定探头并旋转容器进行多周向位置采集，累计获得无油 482 组、有油 461 组，每组 220 点；另完成多距离数据采集。')
add_bullet('编写 DAT 组数统计、波形对比、基础特征分析和自动判断程序，能够完成批量数据解析和可视化。')
add_bullet('完成 STM32F103 离线初步判定链路，有油位置显示红灯、无油位置显示绿灯。')
add_body('这些工作证明硬件链路、数据采集和有限场景下的空气/硅油初步可分性已经建立，是后续开题研究的前期基础。但现有样本仍缺少严格液位真值、跨日与重新安装独立测试、系统角度记录和三介质数据，不能把当前红/绿灯结果表述为已经达到工程准确率。')
add_heading('4.5 尚需补足的关键环节',2)
add_bullet('标准化测量链标定：点序号—等效距离、预热漂移、重复安装和波束足迹。')
add_bullet('可观察液位真值的轴向扫描平台及正反向重复扫描。')
add_bullet('水单介质、三层结构和硅油—水界面数据。')
add_bullet('同材质外壳厚度、站距和入射角系统试验。')
add_bullet('按容器/日期/安装批次划分的独立数据集与冻结盲测集。')
add_bullet('退役真实终端上的工程适用性分析。')
add_heading('4.6 风险控制与备选方案',2)
add_table(['风险','控制或备选方案'],[
['空气与硅油差异不足','优化站距和法向角，使用参考回波与多次回波比值，增加多帧统计；先在薄壁平台完成机理，再逐步增加厚度并明确边界。'],
['水吸收过强','将门后能量骤降和有效回波消失作为水类别特征，重点研究最小可识别水层。'],
['真实终端真值难获得','透明模拟终端作为定量结论主体；真实退役终端仅做正常/低油位工程适用性。'],
['因素组合过大','厚度、距离、角度为主因素；温度和表面凝露为验证因素；采用正交或D-optimal缩减组合。'],
['深度模型样本不足','以低维物理特征和传统模型为主，深度模型仅作对照。']], [4.5,13.0])

# 5
add_heading('5 研究进度安排',1)
add_body('进度安排根据已完成工作重新修订，不再把二维/三维机理仿真、基础实物链路和空气/硅油初步采集列为从零开始的未来任务。')
add_table(['时间','主要任务','阶段成果'],[
['2026.07—2026.09','整理前期仿真与实物数据；完成探头输出、预热漂移、点序号—距离与波束足迹标定；设计标准化平台','前期工作报告、标定方案、平台图纸'],
['2026.10—2026.12','加工可观察真值的模拟终端与轴向/角度平台；完成目标工作频点参数校准与空气/硅油规范数据','平台验收、参数范围、空气/硅油基准数据'],
['2027.01—2027.04','完成水单介质、空气—硅油与三层数据采集；建立分组数据集和介质识别模型','三介质数据集、识别模型、阶段论文初稿'],
['2027.05—2027.07','完成空气—硅油及硅油—水界面扫描、概率变点定位和不确定度评价','界面定位算法与精度结果'],
['2027.08—2027.11','开展外壳厚度、站距、角度及有限温度/凝露试验','影响规律、可靠检测域、SCI论文初稿'],
['2027.12—2028.02','完成STM32轻量化与退役真实终端工程适用性验证','离线原型、PC/MCU一致性、验证报告'],
['2028.03—2028.06','补充关键实验，完成论文投稿、学位论文撰写与答辩','论文投稿、学位论文和答辩材料']], [3.2,8.5,5.8])

# 6
add_heading('6 预期成果',1)
for s in ['形成亚太赫兹波在电缆充油终端局部多层介质中的传播与回波机理模型，并建立目标工作频点实测校准方法。','形成空气、硅油和水三介质回波数据集、数据字典和防止数据泄漏的分组验证流程。','形成空气—硅油和硅油—水双界面定位方法、误差预算与最小可识别水层结果。','形成外壳厚度、探头站距和入射角的可靠检测域、不可测窗口与低置信拒识规则。','完成 RS485—STM32 离线检测流程与退役真实终端工程适用性分析。','力争发表 SCI 论文 1 篇以上，完成相关软件著作权或专利，并完成硕士学位论文。']:
    add_bullet(s)

# 7
add_heading('7 拟定学位论文目录',1)
add_body('学位论文保持六章结构，以“能够识别什么—界面定位多准—哪些条件下仍准确”为主线。THz-TDS、仿真和嵌入式均服务于核心检测问题，不扩展为相互独立的并行主线。')
outline=[
('第1章 绪论',['1.1 研究背景与意义','1.2 电缆充油终端结构、故障形式及检测需求','1.3 液位检测与终端内部状态检测研究现状','1.4 亚太赫兹频段与侧壁检测路线选择依据','1.5 研究目标、研究内容与技术路线','1.6 本章小结']),
('第2章 亚太赫兹侧壁回波检测原理与系统标定',['2.1 终端局部多层介质结构与侧壁检测模型','2.2 多层介质反射、传输与多次回波机理','2.3 空气、硅油和水的介电响应差异','2.4 材料频谱趋势与目标工作频点参数校准','2.5 探头输出、点序号—距离与波束足迹标定','2.6 本章小结']),
('第3章 亚太赫兹侧壁回波仿真与检测方案',['3.1 前期仿真模型与结果复核','3.2 一维传输矩阵与目标频点等效模型','3.3 二维局部曲面与有限三维验证','3.4 模拟终端、扫描机构与数据采集系统','3.5 真值体系、实验设计与数据集构建','3.6 本章小结']),
('第4章 空气—硅油—水介质识别与分层界面定位',['4.1 回波数据预处理与质量控制','4.2 物理门控特征构建与筛选','4.3 空气、硅油和水介质识别','4.4 空气—硅油界面定位','4.5 硅油—水界面定位与最小水层厚度','4.6 识别与定位不确定度评价','4.7 本章小结']),
('第5章 影响因素、适用边界与工程适用性',['5.1 基准条件下的识别与定位性能','5.2 外壳厚度对检测性能的影响','5.3 探头站距和入射角的影响','5.4 温度与表面凝露等环境因素验证','5.5 补偿方法、可靠检测域与低置信拒识','5.6 轻量模型与STM32离线检测','5.7 退役实际终端工程适用性分析','5.8 本章小结']),
('第6章 结论与展望',['6.1 主要研究结论','6.2 论文主要创新点','6.3 研究不足','6.4 后续研究展望'])]
for ch,subs in outline:
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(7); r=p.add_run(ch); r.bold=True; r._element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
    for s in subs:
        p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(0.75); p.paragraph_format.space_after=Pt(2); p.add_run(s)

# References
add_heading('参考文献',1)
refs=[
'[1] Li S, Cao B, Li J, et al. Review of condition monitoring and defect inspection methods for composited cable terminals[J]. High Voltage, 2023, 8(3): 431-444.',
'[2] Adzman M R, Ahmad M H, Jamil M K M, et al. Diagnosis of MV oil filled cable terminations with X-ray imaging and infrared thermography[J]. American Journal of Applied Sciences, 2007, 4(3): 168-170.',
'[3] 莫润阳, 等. 瓷套式电缆终端油位的超声检测[J]. 西北大学学报(自然科学版), 2015, 45(5).',
'[4] Hong X, Zhang B, Liu Y, et al. Deep-learning-based guided wave detection for liquid-level state in porcelain bushing type terminal[J]. Structural Control and Health Monitoring, 2021, 28(1): e2651.',
'[5] Zhang B, Wei Y J, Liu W Y, et al. A liquid level measurement technique outside a sealed metal container based on ultrasonic impedance and echo energy[J]. Sensors, 2017, 17(1): 185.',
'[6] Liu Y, He X, Zhang T, et al. Liquid-level measurement based on out-of-plane energy of Lamb waves[J]. Applied Acoustics, 2023, 210: 109421.',
'[7] Nakagawa T, Hyodo A, Kogo K, et al. Contactless liquid-level measurement with frequency-modulated millimeter wave through opaque container[J]. IEEE Sensors Journal, 2013, 13(3): 926-933.',
'[8] Lees H, Hayes C, Withayachumnankul W. Terahertz coherence tomography for in-container liquid characterization[J]. IEEE Journal of Selected Topics in Quantum Electronics, 2023, 29(5): 8600408.',
'[9] Wang Z, Guo Y, Ren Z, et al. LiqDetector: Enabling container-independent liquid detection with mmWave signals based on a dual-reflection model[J]. Proc. ACM IMWUT, 2024, 7(4).',
'[10] Guo Y, Wang Z, Ren Z, et al. MultiScanner: Enabling simultaneous detection of multiple liquids with mmWave radar based on a composite reflection model[J]. IEEE Transactions on Mobile Computing, 2025.',
'[11] Cheng L, Xu W, Wang X, et al. Defect detection of cable termination based on the microwave reflection method[J]. IEEE Transactions on Dielectrics and Electrical Insulation, 2024, 31(1): 111-120.',
'[12] Zhai M, Locquet A, Citrin D S. Terahertz nondestructive layer thickness measurement and delamination characterization of GFRP laminates[J]. NDT & E International, 2024, 146: 103170.',
'[13] Nsengiyumva W, Zhong S, Zheng L, et al. Sensing and nondestructive testing applications of terahertz spectroscopy and imaging systems[J]. IEEE Transactions on Instrumentation and Measurement, 2023, 72.',
'[14] Guo C, Xu W, Cai M, et al. Application of terahertz nondestructive testing technology in electrical insulation materials[J]. IEEE Access, 2022, 10: 121547-121560.',
'[15] Lee I S, Lee J W. Nondestructive internal defect detection using a CW-THz imaging system in XLPE for power cable insulation[J]. Applied Sciences, 2020, 10(6): 2055.',
'[16] Li J, et al. Terahertz nondestructive testing method of oil-paper insulation debonding and foreign matter defects[J]. IEEE Transactions on Dielectrics and Electrical Insulation, 2022.',
'[17] Mei H, Jiang H, Yin F, et al. Detection of small defects in composite insulators using terahertz technique and deconvolution method[J]. IEEE Transactions on Instrumentation and Measurement, 2020, 69(10): 8146-8155.',
'[18] Xiong W, Li L, Ren J, et al. Terahertz multiple echoes correction and non-destructive testing based on improved wavelet multi-scale analysis[J]. Sensors, 2022, 22(9): 3477.',
'[19] Yuan H, et al. Moisture transfer characteristics in silicone oil-silicone rubber insulation system considering swelling effect[J]. IEEE Transactions on Dielectrics and Electrical Insulation, 2023, 30(5): 2025-2034.',
'[20] Huang Z, et al. Detection of moisture content of silicone oil in high voltage cable terminal based on terahertz time-domain spectroscopy[C]//2024 Power System and Green Energy Conference. 2024.',
'[21] Ellison W J. Permittivity of pure water over the frequency range 0-25 THz and temperature range 0-100 °C[J]. Journal of Physical and Chemical Reference Data, 2007, 36(1): 1-18.',
'[22] Naito K, Kagawa Y, Utsuno S, et al. Dielectric properties of woven fabric glass fiber reinforced polymer-matrix composites in the THz frequency range[J]. Composites Science and Technology, 2009, 69: 2027-2029.',
'[23] Withayachumnankul W, Naftaly M. Fundamentals of measurement in terahertz time-domain spectroscopy[J]. Journal of Infrared, Millimeter, and Terahertz Waves, 2014, 35: 610-637.',
'[24] Dorney T D, Baraniuk R G, Mittleman D M. Material parameter estimation with terahertz time-domain spectroscopy[J]. JOSA A, 2001, 18(7): 1562-1571.',
'[25] Jepsen P U, Cooke D G, Koch M. Terahertz spectroscopy and imaging—modern techniques and applications[J]. Laser & Photonics Reviews, 2011, 5(1): 124-166.',
'[26] Piotrowsky L, Jaeschke T, Kueppers S, et al. Enabling high accuracy distance measurements with FMCW radar sensors[J]. IEEE Transactions on Microwave Theory and Techniques, 2019.',
'[27] Shi S, et al. High-precision echo processing for FMCW radar liquid level measurement[J]. Measurement, 2021.',
'[28] Xu J, et al. Coherent CZT-based frequency estimation for FMCW radar[J]. IEEE Sensors Journal, 2022.',
'[29] Hegazy A M, Alizadeh M, Samir A, et al. Remote material characterization with complex baseband FMCW radar sensors[J]. Progress In Electromagnetics Research, 2023, 177: 107-126.',
'[30] Rieger K, Erni D, Rueter D. Examination of the liquid volume inside metal tanks using noncontact EMATs from outside[J]. IEEE TUFFC, 2021, 68(4): 1314-1324.',
'[31] Mittleman D M. Twenty years of terahertz imaging[J]. Optics Express, 2018, 26(8): 9417-9431.',
'[32] IEC 60836. Specifications for unused silicone insulating liquids for electrotechnical purposes[S].',
'[33] IEC 62067. Power cables with extruded insulation and their accessories for rated voltages above 150 kV up to 500 kV—Test methods and requirements[S].',
'[34] 国家电网有限公司. 国家电网有限公司十八项电网重大反事故措施[S]. 北京: 中国电力出版社.',
'[35] 周子毅. 高压电缆瓷套式终端内部绝缘油液面检测技术研究[D]. 广州: 华南理工大学, 2018.',
'[36] Pozar D M. Microwave Engineering[M]. 4th ed. Hoboken: Wiley, 2012.'
]
for ref in refs:
    p=doc.add_paragraph(); p.paragraph_format.first_line_indent=Cm(-0.7); p.paragraph_format.left_indent=Cm(0.7); p.paragraph_format.line_spacing=1.2; p.paragraph_format.space_after=Pt(2); p.add_run(ref)

# Appendix
add_heading('附录 A　前期工作与后续任务对应关系',1)
add_table(['工作模块','已完成内容','开题后的主要补充'],[
['理论/仿真','多层传播机理；二维与局部三维瞬态模型；幅值、频率和介电参数扫描','目标频点等效模型；材料参数闭环校准；与实测特征对应'],
['实物链路','模拟管有油/无油测试；RS485采集；STM32红/绿灯判定','标准夹具、液位真值、轴向扫描、跨日和重新安装测试'],
['数据与算法','220点/CFAR解析、波形绘制、基础特征与初判','分组数据集、三介质概率分类、双界面定位、不确定度'],
['影响因素','完成若干分距离采集','同材质外壳厚度、站距、入射角及有限环境验证'],
['工程应用','PC—MCU链路已打通','退役真实终端正常/低油位工程适用性分析']], [3.0,6.6,7.8])

# global paragraph formatting and prevent orphan headings
for p in doc.paragraphs:
    if p.style.name=='Normal':
        p.paragraph_format.line_spacing=1.5
    for r in p.runs:
        r.font.name='Times New Roman'
        if r._element.rPr is None: r._element.get_or_add_rPr()
        r._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

# set core props
doc.core_properties.title='基于亚太赫兹侧壁回波的电缆充油终端空气—硅油—水分层状态识别与界面定位研究'
doc.core_properties.subject='硕士研究生学位论文开题报告'
doc.core_properties.keywords='亚太赫兹, 电缆充油终端, 侧壁回波, 硅油液位, 介质识别, 界面定位'
doc.save(DOCX)
print(DOCX)
