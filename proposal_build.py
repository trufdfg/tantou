from __future__ import annotations

from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

OUT = Path('output')
IMG = OUT / 'images'
OUT.mkdir(exist_ok=True)
IMG.mkdir(exist_ok=True)

FONT_REG = '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc'
FONT_BOLD = '/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc'


def font(size: int, bold: bool = False):
    return ImageFont.truetype(FONT_BOLD if bold else FONT_REG, size=size)


def rounded(draw, xy, radius=24, fill='white', outline='#1f4e79', width=3):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def center_text(draw, box, text, fnt, fill='#12345b', spacing=6):
    x1, y1, x2, y2 = box
    bb = draw.multiline_textbbox((0, 0), text, font=fnt, spacing=spacing, align='center')
    w, h = bb[2] - bb[0], bb[3] - bb[1]
    draw.multiline_text(((x1+x2-w)/2, (y1+y2-h)/2), text, font=fnt, fill=fill, spacing=spacing, align='center')


def arrow(draw, start, end, color='#1f77b4', width=7):
    draw.line([start, end], fill=color, width=width)
    import math
    ang = math.atan2(end[1]-start[1], end[0]-start[0])
    l = 22
    for a in (ang+2.55, ang-2.55):
        p = (end[0]+l*math.cos(a), end[1]+l*math.sin(a))
        draw.line([end, p], fill=color, width=width)


def draw_route(path: Path):
    im = Image.new('RGB', (1800, 1120), 'white')
    d = ImageDraw.Draw(im)
    center_text(d, (0, 25, 1800, 105), '图1  课题总体技术路线', font(48, True), '#102a43')
    titles = ['研究目标与\n问题定义', '材料参数与\n目标频点校准', '传播模型与\n分级仿真', '实验平台与\n规范数据采集', '介质识别与\n界面定位', '影响因素与\n工程适用性']
    items = [
        ['空气/硅油/水可识别性', '双界面定位精度', '厚度/距离/角度边界'],
        ['公开介电数据', 'THz-TDS频谱趋势', '目标工作频点反演校准'],
        ['Fresnel与传输矩阵', '二维局部曲面模型', '有限三维机理验证'],
        ['模拟终端与真值系统', '约122 GHz探头侧壁扫描', 'RS485回波与元数据'],
        ['物理门控与归一化', '低维特征与分类模型', '概率跃迁与变点定位'],
        ['外壳厚度/站距/角度', '温度与表面凝露', 'STM32与退役终端验证'],
    ]
    cols = ['#dceeff', '#d7f3f1', '#dff4e5', '#e9f7e7', '#e6eefc', '#e6f4f1']
    x0, gap, w, y1, y2 = 35, 18, 277, 145, 880
    for i in range(6):
        x1 = x0 + i*(w+gap); x2 = x1+w
        rounded(d, (x1,y1,x2,y2), 30, cols[i], '#1e5f8a', 4)
        d.rounded_rectangle((x1,y1,x2,y1+120), radius=30, fill='#176b9a' if i in (0,4) else '#178f8c' if i in (1,2,5) else '#2d9666')
        center_text(d, (x1+8,y1+6,x2-8,y1+116), f'{i+1}\n{titles[i]}', font(28, True), 'white', 4)
        yy = y1+155
        for s in items[i]:
            d.ellipse((x1+25,yy+8,x1+39,yy+22), fill='#176b9a')
            d.multiline_text((x1+52,yy), s, font=font(23), fill='#17324d', spacing=4)
            yy += 122
        if i < 5:
            arrow(d, (x2+2,(y1+y2)//2), (x2+gap-3,(y1+y2)//2), '#137fa3', 7)
    rounded(d, (80,930,1720,1065), 28, '#f7fbff', '#176b9a', 4)
    center_text(d, (105,945,1695,1050), '形成“材料参数—传播机理—仿真—实验—算法—适用边界”闭环研究体系', font(34, True), '#155a7a')
    im.save(path, quality=95)


def draw_principle(path: Path):
    im = Image.new('RGB', (1800, 1120), 'white')
    d = ImageDraw.Draw(im)
    center_text(d, (0,20,1800,100), '图2  亚太赫兹侧壁检测原理及多层回波路径', font(47, True), '#103b6f')
    # Probe
    rounded(d, (55,390,405,650), 35, '#dceeff', '#204f78', 4)
    center_text(d, (75,410,385,625), '亚太赫兹FMCW探头\n中心工作频率约122 GHz\n侧壁发射与接收', font(28, True), '#163d67')
    # Terminal body
    d.rounded_rectangle((630,130,1280,980), radius=70, fill='#f4f4f4', outline='#454545', width=6)
    d.rounded_rectangle((705,185,1205,925), radius=40, fill='#d9efff', outline='#6b6b6b', width=4)
    # layers
    d.rectangle((730,210,1180,440), fill='#e8f6ff')
    d.rectangle((730,440,1180,735), fill='#ffe7a8')
    d.rectangle((730,735,1180,900), fill='#b8e1ff')
    # center conductor
    d.rectangle((930,120,985,995), fill='#c77929', outline='#7d4315', width=3)
    d.text((810,305), '空气层', font=font(34, True), fill='#174b7a')
    d.text((805,565), '硅油层', font=font(34, True), fill='#6b4a00')
    d.text((825,800), '水层', font=font(34, True), fill='#174b7a')
    d.text((1195,235), '复合外壳', font=font(25), fill='#333333')
    d.text((1195,285), '硅橡胶层', font=font(25), fill='#333333')
    # rays
    colors = ['#1f77b4','#2ca02c','#9467bd','#ff7f0e']
    starts=[(410,470),(410,515),(410,560),(410,605)]
    ends=[(720,355),(720,515),(720,640),(900,780)]
    labels=['外表面回波','内壁回波','多次回波','透射/衰减']
    for c,s,e,l in zip(colors,starts,ends,labels):
        arrow(d,s,e,c,5)
        d.text((430,s[1]-36),l,font=font(23),fill=c)
    # interfaces
    d.line((705,440,1205,440), fill='#e05a28', width=5)
    d.line((705,735,1205,735), fill='#155da4', width=5)
    d.text((1285,420),'空气—硅油界面',font=font(26,True),fill='#e05a28')
    d.text((1285,715),'硅油—水界面',font=font(26,True),fill='#155da4')
    # steps
    boxes=[('步骤1','单点识别局部介质'),('步骤2','沿高度方向轴向扫描'),('步骤3','由类别概率跃迁定位界面')]
    yy=160
    for k,(a,b) in enumerate(boxes):
        rounded(d,(1370,yy,1745,yy+210),24,'#f6f9ff','#1f4e79',3)
        d.rounded_rectangle((1390,yy+18,1500,yy+70),radius=15,fill='#174e8a')
        center_text(d,(1390,yy+18,1500,yy+70),a,font(22,True),'white')
        center_text(d,(1510,yy+20,1725,yy+190),b,font(25,True),'#183b66')
        if k<2: arrow(d,(1555,yy+215),(1555,yy+255),'#4a78b5',5)
        yy+=260
    rounded(d,(360,1010,1440,1090),20,'#f7fbff','#173f73',3)
    center_text(d,(380,1015,1420,1085),'侧壁回波用于局部介质识别；轴向扫描用于分层界面定位',font(29,True),'#173f73')
    im.save(path, quality=95)


def draw_platform(path: Path):
    im = Image.new('RGB', (1800, 1120), 'white')
    d = ImageDraw.Draw(im)
    center_text(d,(0,20,1800,100),'图3  模拟终端实验平台与回波数据采集系统',font(47,True),'#103b6f')
    panels=[(35,130,600,900),(620,130,1135,900),(1155,130,1765,900)]
    names=['1  实验对象与机械平台','2  探测与通信链路','3  数据集构建与标注']
    for p,n in zip(panels,names):
        rounded(d,p,28,'#f9fcff','#174e8a',4)
        d.rounded_rectangle((p[0]+2,p[1]+2,p[2]-2,p[1]+90),radius=26,fill='#174e8a')
        center_text(d,(p[0]+15,p[1]+10,p[2]-15,p[1]+83),n,font(29,True),'white')
    # panel 1 terminal
    d.ellipse((120,220,420,310),fill='#f1f7fb',outline='#2f5b7c',width=4)
    d.rectangle((120,265,420,670),fill='#eaf7ff',outline='#2f5b7c',width=4)
    d.ellipse((120,625,420,715),fill='#d9eefc',outline='#2f5b7c',width=4)
    d.rectangle((130,430,410,590),fill='#ffe5a3')
    d.rectangle((130,590,410,665),fill='#b5e0ff')
    d.text((190,335),'空气',font=font(27,True),fill='#24435f')
    d.text((180,480),'硅油',font=font(27,True),fill='#725000')
    d.text((200,615),'水',font=font(27,True),fill='#174b7a')
    d.line((80,230,80,710),fill='#444',width=5)
    for yy in range(250,711,75): d.line((70,yy,90,yy),fill='#444',width=3)
    d.text((45,745),'独立液位真值',font=font(24,True),fill='#17324d')
    d.rectangle((95,800,535,840),fill='#c8d6df',outline='#4d5961',width=3)
    d.rectangle((260,775,350,865),fill='#4f89b2',outline='#24506d',width=3)
    d.text((190,865),'轴向扫描导轨',font=font(24,True),fill='#17324d')
    # panel 2 chain
    blocks=[('亚太赫兹探头\n约122 GHz',220),('RS485接口\n115200 bps',390),('PC完整存储\nSTM32离线验证',575)]
    for txt,yy in blocks:
        rounded(d,(720,yy,1040,yy+120),22,'#e9f4ff','#1f5f94',3)
        center_text(d,(735,yy+8,1025,yy+112),txt,font(26,True),'#174b7a')
    arrow(d,(880,342),(880,385),'#1f77b4',6); arrow(d,(880,512),(880,570),'#1f77b4',6)
    d.text((680,740),'可用观测量',font=font(26,True),fill='#17324d')
    for i,s in enumerate(['220点幅值序列','CFAR阈值序列','Pos/Dis距离信息']):
        d.text((720,785+i*42),'• '+s,font=font(23),fill='#17324d')
    # panel 3 cards
    cards=[('单介质基准数据','空气 / 硅油 / 水'),('两层数据','空气—硅油'),('三层数据','空气—硅油—水'),('影响因素数据','厚度 / 距离 / 角度 / 温度')]
    yy=225
    for a,b in cards:
        rounded(d,(1235,yy,1685,yy+125),20,'#eef9f1','#34865c',3)
        d.text((1260,yy+20),a,font=font(25,True),fill='#245b40')
        d.text((1260,yy+70),b,font=font(22),fill='#345')
        yy+=145
    d.text((1240,825),'同步元数据：容器、日期、安装批次、\n厚度、站距、角度、温湿度、真值',font=font(22),fill='#17324d',spacing=8)
    rounded(d,(95,950,1705,1075),25,'#f7fbff','#174e8a',4)
    center_text(d,(115,960,1685,1065),'实验平台  →  探测链路  →  数据采集  →  分组数据集',font(34,True),'#174e8a')
    im.save(path, quality=95)


def draw_algorithm(path: Path):
    im = Image.new('RGB', (1800, 1120), 'white')
    d=ImageDraw.Draw(im)
    center_text(d,(0,20,1800,100),'图4  回波处理、三介质识别与界面定位流程',font(47,True),'#102a43')
    titles=['原始回波输入','数据预处理','特征构建','三介质识别','界面定位','结果评价与部署']
    bodies=[
        ['220点幅值序列','CFAR序列','Pos/Dis信息'],
        ['帧完整性检查','背景扣除','外表面峰对齐','门控与归一化'],
        ['峰值比/能量比','多次回波衰减','相关与统计特征','质量分数'],
        ['逻辑回归/LDA','SVM/随机森林','输出三类概率','域外拒识'],
        ['概率交点','顺序约束','变点检测','双界面位置'],
        ['F1/MAE/RMSE','厚度-距离-角度边界','不确定度','STM32一致性'],
    ]
    colors=['#dfeeff','#e1f6f2','#e7f5df','#fff1d8','#ffe4d5','#e4eefc']
    x0,gap,w,y1,y2=30,18,277,145,900
    for i in range(6):
        x1=x0+i*(w+gap); x2=x1+w
        rounded(d,(x1,y1,x2,y2),28,colors[i],'#1d5f82',4)
        d.ellipse((x1+w/2-35,y1-35,x1+w/2+35,y1+35),fill='#1f77b4' if i in (0,5) else '#168b83' if i in (1,2) else '#f39c12' if i==3 else '#e66a1e')
        center_text(d,(x1+w/2-35,y1-35,x1+w/2+35,y1+35),str(i+1),font(29,True),'white')
        center_text(d,(x1+15,y1+55,x2-15,y1+140),titles[i],font(28,True),'#17324d')
        yy=y1+185
        for s in bodies[i]:
            d.ellipse((x1+25,yy+8,x1+39,yy+22),fill='#1d6f94')
            d.multiline_text((x1+52,yy),s,font=font(22),fill='#17324d',spacing=3)
            yy+=95
        if i<5: arrow(d,(x2+2,530),(x2+gap-2,530),'#506070',6)
    rounded(d,(130,960,1670,1070),25,'#f7fbff','#174e8a',4)
    center_text(d,(150,970,1650,1060),'从单点回波识别到轴向双界面定位的完整算法链',font(35,True),'#174e8a')
    im.save(path,quality=95)


def draw_progress(path: Path):
    im=Image.new('RGB',(1800,1050),'white'); d=ImageDraw.Draw(im)
    center_text(d,(0,20,1800,100),'图5  前期工作基础与后续研究衔接',font(47,True),'#103b6f')
    cols=[(40,145,570,900),(635,145,1165,900),(1230,145,1760,900)]
    headers=[('已完成','#2d8a5b'),('正在完善','#d78a00'),('后续重点','#1f5f99')]
    tasks=[
        ['侧壁检测机理与多层介质理论梳理','COMSOL二维/三维瞬态模型','入射幅值、频率及硅油参数扫描','RS485通信与220点幅值/CFAR采集','玻纤-硅橡胶模拟终端有油/无油试验','STM32红/绿灯离线判定链路'],
        ['将仿真模型收敛为目标工作频点局部模型','建立点序号—距离与波束足迹标定','规范跨日、跨容器、跨安装数据集','由阈值判别升级为物理特征与概率分类'],
        ['空气—硅油界面轴向定位与精度评价','水介质识别及硅油—水界面拓展','外壳厚度、站距和入射角适用边界','温度/凝露异常与低置信拒识','退役实际终端工程适用性验证'],
    ]
    for idx,(p,(h,c),ls) in enumerate(zip(cols,headers,tasks)):
        rounded(d,p,30,'#fbfdff',c,4)
        d.rounded_rectangle((p[0]+2,p[1]+2,p[2]-2,p[1]+95),radius=28,fill=c)
        center_text(d,(p[0]+10,p[1]+8,p[2]-10,p[1]+88),h,font(32,True),'white')
        yy=p[1]+135
        for s in ls:
            d.ellipse((p[0]+30,yy+7,p[0]+46,yy+23),fill=c)
            d.multiline_text((p[0]+62,yy),s,font=font(23),fill='#17324d',spacing=4)
            yy+=95
        if idx<2: arrow(d,(p[2]+10,520),(cols[idx+1][0]-10,520),'#6b7a88',7)
    rounded(d,(260,935,1540,1020),20,'#f3f8ff','#174e8a',3)
    center_text(d,(280,940,1520,1015),'前期结果证明“可测”；后续研究回答“多准、何时可靠、如何迁移”',font(31,True),'#174e8a')
    im.save(path,quality=95)


def set_run_font(run, east='宋体', latin='Times New Roman', size=12, bold=None, color=None):
    run.font.name = latin
    run._element.rPr.rFonts.set(qn('w:eastAsia'), east)
    run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if color: run.font.color.rgb = RGBColor(*color)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)


def set_cell_text(cell, text, bold=False, size=10.5, color=(0,0,0), align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text=''
    p=cell.paragraphs[0]; p.alignment=align
    r=p.add_run(text); set_run_font(r,size=size,bold=bold,color=color)
    cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader'); tblHeader.set(qn('w:val'), 'true'); trPr.append(tblHeader)


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar=OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=tcMar.find(qn(f'w:{m}'))
        if node is None: node=OxmlElement(f'w:{m}'); tcMar.append(node)
        node.set(qn('w:w'),str(v)); node.set(qn('w:type'),'dxa')


def set_table_borders(table, color='7F8C8D', size='6'):
    tblPr=table._tbl.tblPr
    borders=tblPr.first_child_found_in('w:tblBorders')
    if borders is None: borders=OxmlElement('w:tblBorders'); tblPr.append(borders)
    for edge in ('top','left','bottom','right','insideH','insideV'):
        tag=OxmlElement(f'w:{edge}'); tag.set(qn('w:val'),'single'); tag.set(qn('w:sz'),size); tag.set(qn('w:color'),color); borders.append(tag)


def add_page_number(paragraph):
    paragraph.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run=paragraph.add_run()
    fldChar1=OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'),'begin')
    instrText=OxmlElement('w:instrText'); instrText.set(qn('xml:space'),'preserve'); instrText.text=' PAGE '
    fldChar2=OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'),'end')
    run._r.extend([fldChar1,instrText,fldChar2])


def add_toc(paragraph):
    run=paragraph.add_run()
    begin=OxmlElement('w:fldChar'); begin.set(qn('w:fldCharType'),'begin')
    instr=OxmlElement('w:instrText'); instr.set(qn('xml:space'),'preserve'); instr.text='TOC \\o "1-3" \\h \\z \\u'
    separate=OxmlElement('w:fldChar'); separate.set(qn('w:fldCharType'),'separate')
    text=OxmlElement('w:t'); text.text='打开Word后右键更新目录'
    separate.append(text)
    end=OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'),'end')
    run._r.extend([begin,instr,separate,end])


def add_para(doc, text='', style=None, bold_lead=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True, size=12, space_after=4):
    p=doc.add_paragraph(style=style); p.alignment=align
    p.paragraph_format.line_spacing=1.5
    p.paragraph_format.space_after=Pt(space_after)
    if first_line: p.paragraph_format.first_line_indent=Cm(0.74)
    if bold_lead and text.startswith(bold_lead):
        r=p.add_run(bold_lead); set_run_font(r,size=size,bold=True)
        r=p.add_run(text[len(bold_lead):]); set_run_font(r,size=size)
    else:
        r=p.add_run(text); set_run_font(r,size=size)
    return p


def add_bullet(doc, text, level=0):
    p=doc.add_paragraph(style='List Bullet' if level==0 else 'List Bullet 2')
    p.paragraph_format.line_spacing=1.4; p.paragraph_format.space_after=Pt(2)
    r=p.add_run(text); set_run_font(r,size=11.5)
    return p


def add_caption(doc, text):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(8)
    r=p.add_run(text); set_run_font(r,size=10.5)


def add_figure(doc, path, caption, width=16.0):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width))
    add_caption(doc, caption)


def add_table(doc, headers, rows, widths=None, font_size=9.5):
    table=doc.add_table(rows=1, cols=len(headers)); table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.style='Table Grid'
    set_table_borders(table)
    hdr=table.rows[0]; set_repeat_table_header(hdr)
    for i,h in enumerate(headers):
        shade_cell(hdr.cells[i],'D9EAF7'); set_cell_text(hdr.cells[i],h,bold=True,size=font_size)
    for row in rows:
        cells=table.add_row().cells
        for i,val in enumerate(row):
            set_cell_text(cells[i],str(val),size=font_size,align=WD_ALIGN_PARAGRAPH.LEFT if i>0 else WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_margins(cells[i])
    if widths:
        for row in table.rows:
            for i,w in enumerate(widths): row.cells[i].width=Cm(w)
    doc.add_paragraph()
    return table


def build_document():
    fig1=IMG/'fig1_route.png'; fig2=IMG/'fig2_principle.png'; fig3=IMG/'fig3_platform.png'; fig4=IMG/'fig4_algorithm.png'; fig5=IMG/'fig5_progress.png'
    draw_route(fig1); draw_principle(fig2); draw_platform(fig3); draw_algorithm(fig4); draw_progress(fig5)

    doc=Document()
    sec=doc.sections[0]
    sec.page_width=Cm(21); sec.page_height=Cm(29.7); sec.top_margin=Cm(2.4); sec.bottom_margin=Cm(2.2); sec.left_margin=Cm(2.6); sec.right_margin=Cm(2.4)
    # styles
    styles=doc.styles
    normal=styles['Normal']; normal.font.name='Times New Roman'; normal._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体'); normal.font.size=Pt(12)
    for name,size,bold,color in [('Title',22,True,(14,52,90)),('Heading 1',16,True,(15,59,102)),('Heading 2',14,True,(20,74,120)),('Heading 3',12.5,True,(35,79,117))]:
        st=styles[name]; st.font.name='Times New Roman'; st._element.rPr.rFonts.set(qn('w:eastAsia'),'黑体'); st.font.size=Pt(size); st.font.bold=bold; st.font.color.rgb=RGBColor(*color)
        st.paragraph_format.space_before=Pt(10); st.paragraph_format.space_after=Pt(6)
    # update fields
    settings=doc.settings._element
    upd=OxmlElement('w:updateFields'); upd.set(qn('w:val'),'true'); settings.append(upd)
    # footer
    for s in doc.sections: add_page_number(s.footer.paragraphs[0])

    # cover
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(20)
    r=p.add_run('硕士研究生学位论文开题报告'); set_run_font(r,east='黑体',size=20,bold=True,color=(15,59,102))
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(32)
    r=p.add_run('基于亚太赫兹侧壁回波的电缆充油终端\n空气—硅油—水分层状态识别与界面定位研究'); set_run_font(r,east='黑体',size=23,bold=True,color=(10,45,80))
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(18)
    r=p.add_run('Sub-Terahertz Sidewall-Echo-Based Identification and Interface Localization\nof Air–Silicone Oil–Water Stratification in Oil-Filled Cable Terminations'); set_run_font(r,size=13,bold=False,color=(70,70,70))
    doc.add_paragraph('\n')
    info=[('学　　院','____________________________'),('专　　业','电气工程'),('研究方向','高电压与绝缘技术 / 电气设备状态检测'),('研究生','____________________________'),('指导教师','____________________________'),('日　　期','2026 年 7 月')]
    t=doc.add_table(rows=len(info),cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,(a,b) in enumerate(info):
        set_cell_text(t.cells[i*2] if False else t.rows[i].cells[0],a,bold=True,size=12,align=WD_ALIGN_PARAGRAPH.RIGHT)
        set_cell_text(t.rows[i].cells[1],b,size=12,align=WD_ALIGN_PARAGRAPH.LEFT)
        for c in t.rows[i].cells:
            c._tc.get_or_add_tcPr().append(OxmlElement('w:tcBorders'))
    doc.add_paragraph('\n')
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('文献检索、研究方案与前期工作整理截至 2026 年 7 月'); set_run_font(r,size=10.5,color=(90,90,90))
    doc.add_page_break()

    # Abstract
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('摘　要'); set_run_font(r,east='黑体',size=18,bold=True,color=(15,59,102))
    abstract=(
        '电缆充油终端是高压电缆系统中电场、热场和机械应力高度集中的关键附件。硅油泄漏、油位下降、密封失效及水分侵入会改变终端内部绝缘介质分布和电场边界，若不能及时发现，可能进一步诱发局部放电、异常发热、沿面放电乃至绝缘击穿。现有红外、超声、射线及常规电气诊断方法各有适用条件，但在不拆解、不开孔、不接触内部介质的条件下，从复合绝缘侧壁直接识别空气、硅油和水，并定量定位分层界面，仍缺少系统研究。为此，本课题以亚太赫兹侧壁回波为信息载体，研究电缆充油终端内部空气—硅油—水介质状态的可观测性、可识别性和分层界面的可定位性。实验中选用中心工作频率约122 GHz的线性调频连续波亚太赫兹探头，沿终端轴向进行侧壁扫描，重点回答“能够测到什么、界面能够定位多准、在何种结构与安装条件下仍然准确”三个核心问题。'
    )
    add_para(doc,abstract)
    add_para(doc,'研究首先通过公开介电数据、THz-TDS频谱趋势和目标工作频点实测反演，建立玻纤复合外壳、硅橡胶、空气、硅油和水的等效电磁参数范围；在此基础上构建Fresnel反射、传输矩阵与多次回波模型，并采用“一维解析—二维局部曲面—有限三维验证”的分级仿真策略，揭示外表面、层间、内壁及后续多次回波中介质信息的形成机制。实验上构建可观察真值、可更换等效壁厚、可调站距和角度的模拟终端及轴向扫描平台，形成单介质、空气—硅油两层和空气—硅油—水三层数据集。通过外表面参考回波对齐、背景扣除、物理门控、回波比值与多次回波衰减等特征，实现三介质识别；进一步利用类别概率跃迁、物理顺序约束和变点拟合完成空气—硅油与硅油—水界面定位。')
    add_para(doc,'前期已完成侧壁传播机理分析、COMSOL二维与局部三维瞬态模型、入射幅值/频率/硅油参数扫描、RS485通信与220点幅值—CFAR采集、玻纤—硅橡胶模拟终端有油/无油试验以及STM32红绿灯离线判定链路。上述工作证明了“可测”和硬件链路可运行，但现有仿真频率、网格和观测量尚未与实际探头完全对应，现有实物数据也缺少统一真值、轴向定位和跨工况验证。因此，后续研究将重点完成目标工作频点校准、规范数据集、双界面定位精度、厚度—距离—角度可靠检测域及实际终端工程适用性评价。')
    p=doc.add_paragraph(); r=p.add_run('关键词：'); set_run_font(r,size=12,bold=True); r=p.add_run('电缆充油终端；亚太赫兹；FMCW雷达；侧壁回波；硅油液位；介质识别；界面定位'); set_run_font(r,size=12)
    doc.add_page_break()

    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('目　录'); set_run_font(r,east='黑体',size=18,bold=True,color=(15,59,102)); add_toc(doc.add_paragraph())
    add_para(doc,'图件说明：图1—图5均根据本课题内容独立设计提示词，经image-2生成概念稿后，为保证Word打印清晰度和中文标注准确性进行高清矢量化重绘并嵌入文档。',first_line=False,size=10.5,align=WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_page_break()

    # 1
    doc.add_heading('1 选题背景与研究意义',level=1)
    doc.add_heading('1.1 研究背景',level=2)
    add_para(doc,'随着高压、超高压交联聚乙烯电缆线路持续增长，电缆附件已成为影响输电系统可靠性的薄弱环节之一。终端内部由应力锥、绝缘油、硅橡胶、玻纤或瓷套外壳、密封组件以及贯通电缆等多种材料构成，局部结构复杂，电场、热场和机械应力耦合显著。综合终端状态监测研究表明，制造缺陷、安装偏差、密封老化、材料受潮和运行热应力均可能演化为局部放电、异常发热和绝缘击穿，因此需要在故障形成早期获取内部状态信息[1-3]。')
    add_para(doc,'充油终端依靠硅油填充内部空腔、均化电场并改善散热。密封失效或机械损伤会造成硅油缓慢泄漏，使液面下降并形成上部空气区；外界水分则可能通过密封薄弱部位侵入并在重力作用下沉积于下部，形成“空气—硅油—水”分层结构。空气区的绝缘与散热条件不同于硅油区，水的极性和高介电损耗又会显著改变局部电场及介质损耗。硅油—硅橡胶两相绝缘中的水分迁移还受初始含水量、温度与溶胀效应影响，说明水分异常不仅是单纯的液位问题，也是终端绝缘退化的重要表征[19-20]。')
    add_para(doc,'工程现场通常希望在不停电或少停电、不拆解、不在外壳开孔的条件下判断油位是否正常。然而，终端外壳多为不透明复合绝缘结构，内部空间又被贯通电缆和应力控制部件占据，传统从顶部测量液面的方案难以直接应用。本课题因此采用侧壁检测：在每一轴向高度利用亚太赫兹回波判断探头正对位置的内部介质类别，再由类别沿高度方向的跃迁位置确定分层界面。该路线将“径向回波介质识别”和“轴向扫描界面定位”解耦，可绕开中心电缆遮挡和顶部安装空间受限问题。')
    doc.add_heading('1.2 电缆充油终端故障形式及检测需求',level=2)
    add_para(doc,'围绕本课题，终端内部异常主要包括：密封老化、法兰松动或结构裂纹导致的硅油泄漏与液面下降；雨水、凝露或检修过程引入的水分侵入和底部积水；安装偏心、应力锥位置偏差及内部气隙造成的局部电场畸变；绝缘材料老化、局部过热和表面污染造成的介电参数漂移。检测方法应同时满足以下需求：')
    for s in ['非侵入性：不打开终端、不接触内部硅油，优先从非金属侧壁完成检测。','直接性：检测量与内部空气、硅油或水的介电状态直接相关，而不是仅依赖负荷与热平衡形成的间接温升。','定量性：不仅给出“有油/无油”，还应给出空气—硅油和硅油—水界面位置、误差与置信区间。','鲁棒性：外壳厚度、材料批次、探头距离、入射角度、温度、湿度及表面状态变化时，仍能解释性能变化并给出失效边界。','工程可部署性：算法能够压缩为少量门控特征和轻量分类器，并可在RS485—STM32平台上离线运行。']:
        add_bullet(doc,s)
    doc.add_heading('1.3 检测方法比较及亚太赫兹频段选择依据',level=2)
    add_para(doc,'红外热成像、X射线、超声与电磁反射是当前可用于终端内部状态检测的主要技术方向。红外适合发现运行中异常发热，但温度场受负荷、电流、风速、日照和材料导热性能影响；X射线能够直观显示内部结构，但设备成本、辐射防护和现场布置限制了常态化应用；超声具有成熟的界面反射理论，已有研究通过多次回波衰减或导波特征识别终端油位，但其结果易受耦合层、压力、曲面接触和温度引起的声速变化影响[2-6]。')
    add_para(doc,'亚太赫兹波位于微波与光子学太赫兹之间，兼具毫米波电子学系统的紧凑性与太赫兹波对介电参数差异的敏感性。相较较低频毫米波，其在相同口径下能够形成更小的照射区域，有利于沿终端高度获取局部介质状态；相较更高频电子式太赫兹或激光驱动THz-TDS，其在传播损耗、器件成熟度、体积、成本和现场部署方面更具工程优势。需要强调的是，FMCW名义径向距离分辨率主要由有效扫频带宽决定，载频升高并不自动意味着液位定位精度提高。本课题选择亚太赫兹频段，主要基于穿透能力、介质反射敏感性、波束尺寸和工程实现之间的综合折中。实验设备在后文具体说明为中心工作频率约122 GHz的FMCW亚太赫兹探头。')
    add_table(doc,['检测方法','检测信息','主要优势','主要局限','本课题定位'],[
        ['红外热成像','表面温度场','非接触、巡检成熟','依赖负荷、风速、日照和热传导；早期低油位可能无明显温差','文献比较'],
        ['X射线/数字射线','内部密度与结构投影','结构直观、空间分辨率较高','存在电离辐射、布置和成本限制','高可信参照'],
        ['接触式超声/导波','声阻抗、壁内回波、导波能量','液—气边界敏感、理论成熟','耦合层、压力、曲率和温度影响显著','重点比较对象'],
        ['顶部雷达液位计','液面飞行时间','工业成熟、直接测距','中心电缆贯通且顶部结构复杂','不采用'],
        ['亚太赫兹侧壁回波','介电边界反射、吸收与多次回波','无耦合剂、无电离辐射、可沿侧壁扫描','受厚度、距离、角度和材料离散性影响','核心路线'],
    ],font_size=8.8)
    doc.add_heading('1.4 研究意义',level=2)
    add_para(doc,'理论意义：建立亚太赫兹波在“空气间隙—玻纤复合外壳—硅橡胶—空气/硅油/水”局部曲面多层介质中的传播模型，揭示外表面、层间、内壁及多次反射分量与介质复介电参数之间的映射关系，并阐明轴向扫描如何将径向回波差异转化为分层界面的空间定位问题。')
    add_para(doc,'方法意义：从单一峰值阈值转向“物理门控—多回波比值—归一化特征—概率分类—变点定位”的联合方法，避免仅靠某一幅值点造成的距离与增益敏感性，并以独立测试、不确定度和适用边界替代只报告同工况准确率的评价方式。')
    add_para(doc,'工程意义：形成可从非金属侧壁开展的无耦合剂、无电离辐射检测流程，为红外、超声和射线等现有方法提供补充，并验证低维特征与轻量模型在RS485—STM32离线检测链路上的实现可能性。')

    # 2
    doc.add_heading('2 国内外研究现状与发展趋势',level=1)
    doc.add_heading('2.1 电缆终端状态监测与缺陷检测研究',level=2)
    add_para(doc,'电缆终端状态检测已形成局部放电、红外、超声、射线、电场测量、微波/太赫兹成像及多源融合等技术体系。Li等对复合电缆终端状态监测与缺陷检测进行了系统综述，指出现有方法普遍存在检测对象结构复杂、环境干扰强、定量评价不足和现场可重复性不高等问题，非接触无损检测、智能特征提取和数字化状态评估是重要发展方向[1]。')
    add_para(doc,'针对终端内部介电不连续，微波反射研究已建立含缺陷多层介质模型，并分析缺陷类型、尺寸、硅橡胶厚度和检测距离的影响[11]。该类研究与本课题在“多层介质反射—厚度/距离敏感性—实验验证”的逻辑上高度一致，但其对象主要是固体绝缘内部缺陷，而本课题关注空气、硅油和水三种宏观介质状态及分层界面。')
    doc.add_heading('2.2 终端与密闭容器液位检测研究',level=2)
    add_para(doc,'超声是外置液位检测的主要路线。瓷套式电缆终端研究发现，油区与空气区的首界面回波差异可能较小，多次回波衰减反而具有更稳定的辨识能力[3]；基于导波、声阻抗和Lamb波能量的方法进一步证明沿壁移动传感器能够判断局部液体负载[4-6]，但耦合剂、曲面压力、壁厚共振、频率选择和重复安装误差仍是主要工程约束。')
    add_para(doc,'毫米波非接触液位研究为本课题提供了直接电磁依据。频率调制毫米波可穿透不透明容器实现液位检测[7]；220—330 GHz扫频反射系统可同时估计容器与内部液体信息[8]；LiqDetector与MultiScanner利用多界面反射比值和复合反射模型削弱容器与位置因素[9-10]。这些方法多面向规则、小型容器或具备复数基带数据的系统，尚未解决电缆终端复合厚壁、中心电缆散射和现场安装偏差问题。')
    doc.add_heading('2.3 太赫兹无损检测与多层介质测量研究',level=2)
    add_para(doc,'太赫兹波对低损耗绝缘材料、多层结构和含水缺陷具有较高敏感性，相关研究已用于XLPE电缆内部缺陷、油纸绝缘脱粘与异物、复合绝缘子小缺陷以及多次回波校正[13-18]。GFRP层合板研究利用THz-TDS与稀疏处理实现了各层厚度和分层缺陷的三维定量定位[12]。这些成果说明玻纤复合材料具有一定太赫兹可穿透性，且层内多次反射既可能造成峰重叠，也能够提供厚度与损耗信息。')
    add_para(doc,'THz-TDS能够由样品与参考信号的幅值和相位获得复折射率、吸收系数和复介电常数，但结果对样品厚度、平行度、时窗、相位展开和Fabry—Perot效应敏感[21-23]。对于本课题实际探头的目标工作频点，若实验室TDS系统在低频端信噪比较低，不应将更高频段参数简单外推，而应采用“公开宽带数据—TDS频谱趋势—目标频点标准样件反演”三级约束方式。')
    doc.add_heading('2.4 回波处理、介质识别与界面定位研究',level=2)
    add_para(doc,'FMCW测距研究通常关注频谱泄漏、有限采样、相位噪声和温漂造成的峰值偏差。但本课题探头当前输出为处理后的220点幅值与CFAR序列，缺少原始复数基带、扫频相位和天线S参数，不能直接照搬依赖复相位的高精度算法。研究需要把探头视为具有固定内部处理链的测量设备，通过标准反射体、平移扫描和参考回波建立“点序号—等效距离—幅值—质量分数”的标定关系。')
    add_para(doc,'液位界面定位的本质是沿高度的一维变点问题。界面附近由于波束同时覆盖两种介质，分类概率通常呈连续过渡，而不是理想阶跃。因此应采用概率交点、逻辑曲线拟合、单调约束或变点检测，并将扫描步长、波束轴向足迹、重复定位误差和界面弯月面纳入不确定度预算。')
    doc.add_heading('2.5 现有研究不足与本课题切入点',level=2)
    for s in ['现有电缆终端油位研究以超声、红外和射线为主，面向复合非金属侧壁的亚太赫兹回波研究不足。','太赫兹电力绝缘研究主要关注固体内部微缺陷、层厚和水分光谱，较少研究“复合外壳—硅橡胶—液体”结构中的局部介质识别。','常规雷达液位研究以顶部测距为主，而电缆终端中心电缆贯通、顶部空间受限，需要侧壁扫描和类别跃迁定位的新范式。','现有商用探头只能输出220点幅值和CFAR序列，必须发展与现有接口相匹配、可解释且可嵌入式实现的方法。','既有研究多在单一工况下报告准确率，缺少外壳厚度、距离、角度变化后的跨工况验证、可靠检测域和低置信拒识。']:
        add_bullet(doc,s)
    add_para(doc,'据此，本课题以“现有探头可观测性”为起点，以空气—硅油识别和界面定位为必做核心，以水介质识别为拓展，以外壳厚度、距离和角度适用边界为主要评价，以透明模拟件和退役实际终端分级验证为落点。')

    # 3
    doc.add_heading('3 研究计划',level=1)
    add_para(doc,'本研究计划参照博士开题报告的组织逻辑，将研究目标、研究内容、研究方案、关键问题和创新点纳入同一闭环。各研究环节遵循“参数约束模型、模型指导实验、实验校正模型、识别支撑定位、影响因素界定边界、工程验证检验迁移能力”的因果关系，避免材料测试、仿真、算法和嵌入式相互割裂。')
    doc.add_heading('3.1 研究目标、科学问题与考核指标',level=2)
    doc.add_heading('3.1.1 总体研究目标',level=3)
    add_para(doc,'面向电缆充油终端不透明非金属侧壁和内部贯通电缆结构，建立亚太赫兹侧壁回波检测方法，回答“现有探头能够稳定观察哪些内部信息、空气—硅油—水能否可靠区分、双界面能够定位到何种精度、在何种外壳与安装/环境条件下结果仍可信”四个核心问题。通过材料参数约束、传播模型、分级仿真、模拟终端实验、多因素鲁棒性分析、嵌入式部署和实际终端工程验证，形成可解释、可复现、可量化适用边界的检测体系。')
    doc.add_heading('3.1.2 核心科学与技术问题',level=3)
    add_table(doc,['编号','问题','具体内涵'],[
        ['Q1','可观测性','亚太赫兹波穿过空气间隙、复合外壳和硅橡胶后，哪些门控区仍含有内部空气、硅油和水的稳定差异？'],
        ['Q2','可识别性','如何构造对系统增益、站距和局部曲率不敏感、同时适合单片机计算的低维物理特征？'],
        ['Q3','可定位性','如何把沿轴向离散测点的介质概率转换为双界面位置，并分离扫描步长、波束足迹和模型误差？'],
        ['Q4','鲁棒性与边界','厚度、距离、角度、温度和表面凝露如何影响信号，何时模型性能超出可靠范围？'],
        ['Q5','模型迁移','模拟终端模型如何迁移到真实退役终端，并避免少量校准样本污染测试集？'],
    ],font_size=9.3)
    doc.add_heading('3.1.3 预期技术指标',level=3)
    add_table(doc,['评价对象','阶段目标','评价方式'],[
        ['空气/硅油识别','基准条件宏平均F1≥0.90；跨日/跨安装F1≥0.85','按容器、日期和安装批次分组留出，报告95%置信区间'],
        ['三介质识别','模拟平台宏平均F1≥0.85','水类别单独报告召回率和误报率，不采用帧随机切分'],
        ['空气—硅油界面','模拟终端MAE≤10 mm或不超过一个有效扫描步长','独立盲测，报告偏差、RMSE、最大误差与重复性'],
        ['硅油—水界面','模拟平台MAE≤15 mm','按水层厚度分组，给出检测概率与最小可识别层厚'],
        ['影响因素边界','给出厚度—距离—角度可靠检测域','以F1、界面MAE、内壁门SNR及低置信比例共同定义'],
        ['嵌入式实现','PC与MCU同一输入判定一致','比较特征数值、分类结果、RAM/Flash和单点运算时间'],
    ],font_size=9.0)
    doc.add_heading('3.2 主要研究内容',level=2)
    add_table(doc,['层级','研究任务','完成要求','研究边界'],[
        ['必做主线','空气/硅油局部介质识别；空气—硅油界面定位','跨日、跨容器、跨重新安装测试；报告F1、MAE及置信区间','对应漏油和低油位，是论文最小完整闭环'],
        ['重点研究','外壳厚度、探头站距和入射角影响','建立主效应、交互趋势、补偿方法和可靠检测域','主要学术增量'],
        ['拓展内容','水介质识别、硅油—水界面定位、最小水层厚度','优先在透明模拟终端完成定量验证','真实终端积水仅在安全且真值明确时开展'],
        ['辅助验证','THz-TDS趋势、有限三维仿真、STM32实现','用于参数约束、机理解释和工程实现验证','不开展全尺寸全参数三维扫频'],
    ],font_size=8.8)
    doc.add_heading('3.2.1 材料频谱表征与目标工作频点参数校准',level=3)
    add_para(doc,'利用实验室THz-TDS系统对玻纤复合材料、硅橡胶、硅油和水开展透射或反射测量，获得频谱趋势和参数范围。对固体样品控制厚度、平行度和表面粗糙度，对液体采用固定有效光程样品池，并设置空气、空池和重复装样参考。若TDS在目标工作频点附近动态范围不足，则不将更高频段结果直接替代，而通过公开宽带模型、标准厚度样件和实际探头回波反演进行闭环校准。')
    doc.add_heading('3.2.2 侧壁多层介质传播与分级仿真',level=3)
    add_para(doc,'将探头正对位置局部简化为空气间隙—玻纤复合外壳—硅橡胶—内部介质的分层结构。一维传输矩阵模型用于快速扫描介电参数、层厚和频率，识别Fabry—Perot增强/相消窗口；二维局部曲面模型引入外壳曲率、斜入射和有限空气间隙；有限三维模型只用于验证贯通电缆与有限波束的影响。现有COMSOL三维瞬态模型作为前期机理证据，但后续必须根据实际探头带宽、观测量和目标频点重新约束。')
    doc.add_heading('3.2.3 空气、硅油和水介质识别',level=3)
    add_para(doc,'在同一外壳、同一位置和同一安装条件下构造空气、硅油和水单介质状态，获取基准回波；随后构建两层与三层结构，沿轴向扫描形成带真值标签的数据。原始数据包括220点回波幅值、CFAR序列、Pos/Dis信息、液位真值、外壳厚度、站距、角度、温湿度、日期和安装批次。预处理采用帧完整性检查、背景扣除、外表面峰对齐、物理门控、幅值归一化和距离补偿；特征优先采用外/内壁峰值比、门控能量比、多次回波衰减率、峰位差、峰宽、相关系数、质心、偏度、峰度和CFAR超越面积。模型优先比较LDA、逻辑回归、SVM和随机森林。')
    doc.add_heading('3.2.4 分层界面定位与不确定度',level=3)
    add_para(doc,'探头沿终端轴向以固定步长扫描，在每个位置输出空气、硅油和水的后验概率。空气—硅油界面由P(air)=P(oil)交点或对数似然比过零位置确定，硅油—水界面同理；界面附近采用重叠扫描，并用逻辑曲线、单调样条或带物理顺序约束的变点模型拟合概率序列。定位误差不直接由径向FMCW距离分辨率决定，而主要取决于轴向扫描步长、波束足迹、机械定位误差、弯月面真值和分类概率过渡宽度。')
    doc.add_heading('3.2.5 影响因素、适用边界与失效判据',level=3)
    add_para(doc,'主要因素设置为外壳等效厚度、探头站距和入射角；温度与表面水膜/凝露作为工程验证因素。采用“单因素确认机理—有限多因素组合—留出工况验证”的策略，每个组合包含独立重新安装和完整轴向扫描。响应量包括内壁门SNR、分类F1、界面MAE和低置信样本比例。适用边界以性能阈值定义：当内壁门无法稳定提取、分类性能低于阈值、界面95%误差上限超过允许范围或质量分数不足时，判定超出可靠检测域，而不是继续强制分类。')
    doc.add_heading('3.2.6 轻量模型与实际终端工程适用性',level=3)
    add_para(doc,'嵌入式部分仅部署已在独立测试中验证的低维特征和轻量模型。STM32完成RS485帧接收、220点解析、质量检查、特征计算、介质判定和置信度输出；连续轴向扫描的界面拟合可先在PC端完成。验证按照“材料件—透明模拟件—等效复合壁模拟终端—退役真实终端”逐级进行。透明模拟终端承担定量结论和严格真值验证；退役真实终端重点验证正常、低油位和未知油位，不在带电或不可拆解设备内人为构造漏油、积水故障。')
    add_figure(doc,fig1,'图1  课题总体技术路线（基于image-2概念稿矢量化重绘）',16.2)
    doc.add_heading('3.3 研究方案与实验设计',level=2)
    doc.add_heading('3.3.1 侧壁检测原理',level=3)
    add_para(doc,'探头从侧壁入射，最先形成外表面强回波；部分能量穿过玻纤外壳和硅橡胶后，在内壁—内部介质边界形成回波，并在多层中产生后续多次反射。空气、硅油和水的复介电参数不同，使反射系数和穿透衰减发生变化。单点测量用于判断局部介质，连续轴向扫描用于寻找类别跃迁位置。')
    add_figure(doc,fig2,'图2  亚太赫兹侧壁检测原理及多层回波路径（基于image-2概念稿矢量化重绘）',16.2)
    doc.add_heading('3.3.2 实验平台与数据采集系统',level=3)
    add_para(doc,'模拟终端由透明内胆、可更换等效外壳、硅橡胶层和中心贯通电缆等效件组成。外壳模块应允许更换厚度而不改变内径和曲率；液位真值由透明内胆刻度和独立位移测量给出。探头安装在轴向导轨和二维角度调节机构上，可控制高度、站距、俯仰角和偏航角。实验探测设备为中心工作频率约122 GHz的FMCW亚太赫兹探头，RS485接口波特率115200 bps，可输出Pos、Dis、220点幅值及CFAR序列。PC端保存完整输出，STM32端用于后期离线一致性验证。')
    add_figure(doc,fig3,'图3  模拟终端实验平台与回波数据采集系统（基于image-2概念稿矢量化重绘）',16.2)
    doc.add_heading('3.3.3 分阶段实验方案',level=3)
    add_table(doc,['阶段','主要任务','关键控制','阶段输出'],[
        ['A 测量链标定','金属板平移、角度扫描、重复安装、开机漂移','固定反射板面积与法向基准；测量轴向/周向波束足迹','点序号—距离关系、重复性和质量门限'],
        ['B 单介质基准','空气、硅油、水重复采集','跨日、跨测点和重新安装；液体温度稳定','可观测门控区和候选特征'],
        ['C 分层定位','空气—硅油及空气—硅油—水轴向扫描','独立真值；正反向扫描；界面附近小步长','概率序列、双界面位置和不确定度'],
        ['D 主要因素','厚度、站距、角度','先单因素后L9或D-optimal；每组合独立安装≥3次','主效应、交互作用和可靠域'],
        ['E 环境验证','温度、表面水膜/凝露、轻度污染','同步记录温度与探头预热时间','温度补偿、异常质量门控'],
        ['F 真实终端','退役终端正常/低油位盲测','测试人员与真值人员分离；测试集冻结','域偏移、少量校准收益和工程适用性'],
    ],font_size=8.5)
    doc.add_heading('3.3.4 数据处理、分类与界面定位',level=3)
    add_para(doc,'算法链与现有探头输出保持一致。原始输入为220点幅值、CFAR序列及位置/距离信息；预处理以帧完整性、背景扣除、外表面峰对齐和门控归一化为主；特征构建优先使用峰值比、能量比、多次回波衰减和波形统计特征；分类输出三类概率；界面定位使用概率跃迁和变点检测；最后报告分类、定位、适用边界和嵌入式资源。')
    add_figure(doc,fig4,'图4  回波处理、三介质识别与界面定位流程（基于image-2概念稿矢量化重绘）',16.2)
    doc.add_heading('3.3.5 数据集划分与统计分析',level=3)
    add_para(doc,'每个DAT文件包含大量相邻帧，这些帧高度相关，不能视为独立样本。实验统计单位至少应为稳定测点窗口，更严格时为一条完整轴向扫描。训练、验证和测试按容器、日期、测点组、重新安装批次或真实终端分组，严禁把同一DAT文件相邻帧随机拆分后报告高准确率。模型评价包括混淆矩阵、平衡准确率、宏平均F1、各类别召回率、Brier分数和概率校准；界面评价包括偏差、MAE、RMSE、最大误差、95%误差范围、重复性和正反向扫描差异。影响因素采用方差分析或线性混合效应模型。')
    doc.add_heading('3.4 拟解决的关键问题与应对措施',level=2)
    add_table(doc,['关键问题','主要风险','应对措施'],[
        ['外表面强回波掩盖内部信息','空气/硅油差异不可见或被固定杂波主导','标准反射板标定；外表面峰对齐；物理门控；背景差分；参考峰比值；利用多次回波能量'],
        ['厚度引起Fabry—Perot增强/相消','不同厚度下同一类别特征方向反转','一维传输矩阵扫描敏感窗口；厚度作为模型输入；分厚度模型或不可测窗口'],
        ['站距和角度变化大于介质差异','跨安装误判和概率失真','刚性夹具；参考回波；距离补偿；角度增强训练；低置信拒识'],
        ['界面附近波束混合','概率过渡宽、定位点依赖扫描步长','实测波束足迹；重叠扫描；概率曲线拟合；正反向扫描；置信区间'],
        ['仿真与实际探头不匹配','现有高频脉冲模型不能直接解释实际幅值/CFAR输出','重建目标工作频点局部模型；采用等效参数范围；只比较可观测幅值/能量趋势'],
        ['重复帧数据泄漏','测试准确率虚高','按文件/测点/日期/容器/终端分组；测试集冻结；报告扫描级结果'],
    ],font_size=8.7)
    doc.add_heading('3.5 预期创新点',level=2)
    add_para(doc,'（1）提出适用于中心电缆贯通结构的亚太赫兹侧壁分层检测范式。区别于顶部雷达直接测距，先利用局部径向回波识别空气、硅油和水，再通过轴向扫描的类别跃迁定位分层界面。',first_line=False)
    add_para(doc,'（2）建立与现有幅值/CFAR输出相匹配的“外表面参考—内壁响应—多次回波衰减”联合表征方法，通过物理门控和无量纲比值降低站距、增益和外壳强反射的影响。',first_line=False)
    add_para(doc,'（3）建立带物理顺序约束和不确定度评价的双界面定位方法，利用三类概率曲线和变点检测统一定位空气—硅油与硅油—水界面，并区分扫描步长、波束足迹和分类误差。',first_line=False)
    add_para(doc,'（4）从单一条件识别扩展到厚度—距离—角度适用边界，通过分组独立测试和混合效应分析给出可靠检测域、不可测窗口和低置信拒识规则。',first_line=False)
    add_para(doc,'上述创新点最终以实验结果和持续文献检索为准，开题阶段不使用“首次”“完全解决”等绝对表述。')

    # 4 preliminary
    doc.add_heading('4 可行性论证与前期工作基础',level=1)
    doc.add_heading('4.1 理论可行性',level=2)
    add_para(doc,'空气、硅油和水具有不同的复介电参数，玻纤复合材料和硅橡胶属于可被毫米波/亚太赫兹波一定程度穿透的非导电绝缘材料。Fresnel反射、传输矩阵和多次回波理论可用于描述局部多层结构；毫米波穿容器液位测量、太赫兹容器内液体表征和GFRP多层检测已提供直接方法依据[7-18]。')
    doc.add_heading('4.2 设备与技术可行性',level=2)
    add_para(doc,'现有条件包括亚太赫兹FMCW探头、RS485数据接口、220点幅值与CFAR采集程序、STM32F103开发板、玻纤—硅橡胶模拟终端、学院THz-TDS设备以及COMSOL仿真基础。上述条件覆盖“探头—采集—算法—嵌入式”的基本链路，后续主要新增高精度轴向平台、可更换壁厚结构、角度调节机构和独立真值系统，设备投入与加工难度可控。')
    doc.add_heading('4.3 已完成的前期工作',level=2)
    add_para(doc,'（1）已完成侧壁检测理论与仿真流程。基于麦克斯韦方程、复介电常数、波阻抗和菲涅尔反射关系，建立了空气—硅橡胶—玻纤—空气/硅油多层传播解释框架；在COMSOL中完成二维及局部三维瞬态模型，采用高斯调制脉冲提取油位上方与下方测点的回波响应，并开展入射电场幅值、中心频率和硅油介电参数扫描。',first_line=False)
    add_para(doc,'现有仿真结果表明，非零激励范围内上下测点绝对幅值随入射场近似线性变化，而幅值比在当前模型中保持稳定（约1.412），说明相对幅值、归一化能量和对数能量差更适合作为跨功率、跨增益判据。该结果证明了侧壁回波对空气/硅油路径差异具有可观测性，但由于现有模型采用机理验证频率、粗网格及理想边界，不能直接作为实际探头定量结论。')
    add_para(doc,'（2）已完成探头通信与波形采集。通过RS485指令能够稳定获取Pos、Dis、220点幅值及CFAR序列，已编写批量解析、波形绘制、基础特征提取和自动判定程序；已在多个站距和多个周向位置采集有油/无油数据。',first_line=False)
    add_para(doc,'（3）已完成模拟终端初步实物验证。现有模拟件由玻璃纤维与硅橡胶管构成，内部填充硅油；探头从侧壁分别测量有油区和无油区，获得明显不同的回波波形。探头已接入STM32F103平台，初步实现有油区红灯、无油区绿灯的离线状态指示，证明硬件链路、通信协议和嵌入式判定流程能够运行。该结果仍属于有限场景可分性观察，尚不能代表界面定位精度和跨条件鲁棒性。')
    add_para(doc,'（4）已形成明确的改进认识。现有数据中同一DAT文件相邻帧高度相关，不能随机拆分训练与测试；现有Dis不能直接等同液位高度，液位必须通过轴向扫描的介质类别变化定位；探头轻微距离与角度变化会显著改变信号，需要通过夹具、参考回波、补偿或适用边界处理。',first_line=False)
    add_figure(doc,fig5,'图5  前期工作基础与后续研究衔接（基于image-2概念稿矢量化重绘）',16.2)
    doc.add_heading('4.4 当前研究缺口与下一步衔接',level=2)
    for s in ['仿真端：现有二维/三维模型与实际探头的中心频率、扫频带宽、天线波束和220点输出未完全对应，需要重建目标工作频点的局部等效模型。','材料端：缺少玻纤、硅橡胶、硅油及水在目标工作频点的可靠参数，应采用公开模型、TDS趋势和标准样件反演校准。','实验端：现有有油/无油测试缺少高精度轴向真值、固定站距/角度夹具、跨日和重新安装数据，需要搭建规范扫描平台。','算法端：现有阈值判别适合初步演示，但论文需要物理门控、多特征模型、分组独立测试、概率校准及低置信拒识。','研究对象：尚未系统采集水介质与双界面数据，应先完成空气—硅油主线，再拓展硅油—水界面和最小水层厚度。']:
        add_bullet(doc,s)
    doc.add_heading('4.5 风险控制与备选方案',level=2)
    add_table(doc,['风险','控制或备选方案'],[
        ['空气与硅油回波差异不足','优化站距和法向角；使用外表面参考与多回波比值；增加多帧统计；先在薄壁平台完成机理，再逐步增加厚度并明确边界。'],
        ['水吸收过强导致后续回波消失','把门后能量骤降和有效回波消失作为水类别特征；重点研究最小可识别水层。'],
        ['真实终端真值难获得','透明模拟终端作为定量结论主体；真实终端仅做正常/低油位工程验证。'],
        ['因素组合数量过大','厚度、距离、角度为主因素；温度和凝露为验证因素；采用L9或D-optimal缩减组合。'],
        ['深度模型样本不足','以低维物理特征和传统模型为主；深度模型仅作对照。'],
    ],font_size=9.0)

    # 5 schedule
    doc.add_heading('5 研究进度安排',level=1)
    add_para(doc,'进度安排以已经完成的工作为起点，不再把理论机理、基础COMSOL模型、RS485通信和有油/无油实物演示写成未来任务。后续时间主要投入规范标定、界面定位、影响因素和工程适用性验证。')
    add_table(doc,['时间','主要任务','阶段成果'],[
        ['已完成（截至2026.07）','完成文献与机理梳理、二维/局部三维仿真、幅值/频率/硅油参数扫描、RS485采集、模拟终端有油/无油测试、STM32状态指示','前期技术报告、仿真模型、采集程序、初步实物链路'],
        ['2026.07—2026.09','完成探头点序号—距离、波束足迹、重复安装和开机漂移标定；重建目标工作频点局部模型','标定报告、仿真基线、平台改造图纸'],
        ['2026.10—2026.12','完成可更换壁厚模拟终端和轴向/角度平台；完成材料样片与TDS趋势测试','平台验收、参数范围、单介质预实验'],
        ['2027.01—2027.04','完成空气/硅油/水基准数据和规范数据集；建立介质识别模型','分组数据集、识别模型、阶段论文'],
        ['2027.05—2027.07','完成空气—硅油和硅油—水轴向扫描、定位与不确定度评价','界面定位算法、精度结果'],
        ['2027.08—2027.11','开展厚度、距离、角度因素试验及有限温度/凝露验证','可靠检测域、补偿与拒识规则、SCI论文初稿'],
        ['2027.12—2028.02','完成STM32轻量实现和退役实际终端工程适用性验证','离线原型、PC/MCU一致性、验证报告'],
        ['2028.03—2028.06','补充关键实验，完成投稿和学位论文撰写','论文投稿、学位论文和答辩材料'],
    ],font_size=8.5)

    doc.add_heading('6 预期成果',level=1)
    for s in ['建立亚太赫兹波在电缆终端局部多层介质中的传播与回波机理模型，并完成目标工作频点等效参数校准。','形成空气、硅油和水介质回波数据集、数据字典及可复现实验流程。','形成空气—硅油和硅油—水双界面定位方法、误差模型和最小可识别液层结果。','形成外壳厚度、站距和入射角的可靠检测域及低置信拒识规则。','完成RS485—STM32轻量离线检测流程和退役实际终端工程适用性分析。','力争发表SCI论文1篇以上，完成相关专利或软件著作权，并完成硕士学位论文。']:
        add_bullet(doc,s)

    doc.add_heading('7 拟定学位论文目录',level=1)
    outline=[
        ('第1章 绪论',['1.1 研究背景与意义','1.2 电缆充油终端结构、故障形式及检测需求','1.3 液位检测及电缆终端内部状态检测研究现状','1.4 不同检测方法比较及亚太赫兹频段选择依据','1.5 研究目标、主要研究内容与拟解决的关键问题','1.6 技术路线']),
        ('第2章 亚太赫兹侧壁回波检测原理与材料参数表征',['2.1 电缆充油终端侧壁检测结构与多层介质模型','2.2 亚太赫兹波检测原理及多层介质传播方式','2.3 空气、硅油和水的电磁响应差异','2.4 基于THz-TDS的材料频谱特性与目标工作频点校准','2.5 沿轴向扫描与分层界面定位原理','2.6 本章小结']),
        ('第3章 亚太赫兹侧壁回波仿真与检测方案设计',['3.1 仿真目标、模型假设与建模方法','3.2 局部多层介质模型及参数设置','3.3 空气、硅油和水三介质识别仿真','3.4 外壳厚度、距离与角度的影响分析','3.5 仿真结果与实验方案的对应关系','3.6 亚太赫兹侧壁检测总体方案','3.7 本章小结']),
        ('第4章 空气—硅油—水分层识别与界面定位',['4.1 模拟电缆充油终端实验平台','4.2 探头、回波数据采集系统及电路搭建','4.3 三介质识别试验设计与数据集构建','4.4 回波数据预处理方法','4.5 回波特征构建与筛选','4.6 空气、硅油和水介质识别','4.7 空气—硅油界面定位方法','4.8 硅油—水界面定位方法','4.9 识别与定位性能评价','4.10 本章小结']),
        ('第5章 影响因素、适用边界与工程适用性',['5.1 基准条件下的识别与定位性能','5.2 外壳厚度对检测性能的影响','5.3 探头距离对检测性能的影响','5.4 入射角度与安装偏差对检测性能的影响','5.5 环境条件、最小液层厚度与空间分辨能力','5.6 检测方法的适用边界与不确定度','5.7 轻量模型与单片机离线检测系统','5.8 退役实际终端工程适用性分析','5.9 与其他检测方法的综合比较','5.10 本章小结']),
        ('第6章 结论与展望',['6.1 主要研究结论','6.2 论文主要创新点','6.3 研究不足','6.4 后续研究展望']),
    ]
    for h,subs in outline:
        p=doc.add_paragraph(); r=p.add_run(h); set_run_font(r,east='黑体',size=12,bold=True,color=(15,59,102))
        for s in subs:
            p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(0.7); r=p.add_run(s); set_run_font(r,size=11.2)

    doc.add_heading('参考文献',level=1)
    refs=[
        '[1] Li S, Cao B, Li J, et al. Review of condition monitoring and defect inspection methods for composited cable terminals[J]. High Voltage, 2023, 8(3): 431-444.',
        '[2] Adzman M R, Ahmad M H, Jamil M K M, et al. Diagnosis of MV oil filled cable terminations with X-ray imaging and infrared thermography[J]. American Journal of Applied Sciences, 2007, 4(3): 168-170.',
        '[3] 莫润阳, 等. 瓷套式电缆终端油位的超声检测[J]. 西北大学学报(自然科学版), 2015, 45(5).',
        '[4] Hong X, Zhang B, Liu Y, et al. Deep-learning-based guided wave detection for liquid-level state in porcelain bushing type terminal[J]. Structural Control and Health Monitoring, 2021, 28(1): e2651.',
        '[5] Zhang B, Wei Y J, Liu W Y, et al. A liquid level measurement technique outside a sealed metal container based on ultrasonic impedance and echo energy[J]. Sensors, 2017, 17(1): 185.',
        '[6] Liu Y, He X, Zhang T, et al. Liquid-level measurement based on out-of-plane energy of Lamb waves[J]. Applied Acoustics, 2023, 210: 109421.',
        '[7] Nakagawa T, Hyodo A, Kogo K, et al. Contactless liquid-level measurement with frequency-modulated millimeter wave through opaque container[J]. IEEE Sensors Journal, 2013, 13(3): 926-933.',
        '[8] Lees H, Hayes C, Withayachumnankul W. Terahertz coherence tomography for in-container liquid characterization[J]. IEEE Journal of Selected Topics in Quantum Electronics, 2023, 29(5): 8600408.',
        '[9] Wang Z, Guo Y, Ren Z, et al. LiqDetector: Enabling container-independent liquid detection with mmWave signals based on a dual-reflection model[J]. Proc. ACM IMWUT, 2024, 7(4).',
        '[10] Guo Y, Wang Z, Ren Z, et al. MultiScanner: Enabling simultaneous detection of multiple liquids with mmWave radar based on a composite reflection model[J]. IEEE Transactions on Mobile Computing, 2025.',
        '[11] Cheng L, Xu W, Wang X, et al. Defect detection of cable termination based on the microwave reflection method[J]. IEEE Transactions on Dielectrics and Electrical Insulation, 2024, 31(1): 111-120.',
        '[12] Zhai M, Locquet A, Citrin D S. Terahertz nondestructive layer thickness measurement and delamination characterization of GFRP laminates[J]. NDT & E International, 2024, 146: 103170.',
        '[13] Nsengiyumva W, Zhong S, Zheng L, et al. Sensing and nondestructive testing applications of terahertz spectroscopy and imaging systems[J]. IEEE Transactions on Instrumentation and Measurement, 2023, 72.',
        '[14] Guo C, Xu W, Cai M, et al. Application of terahertz nondestructive testing technology in electrical insulation materials: A review[J]. IEEE Access, 2022, 10: 121547-121560.',
        '[15] Lee I S, Lee J W. Nondestructive internal defect detection using a CW-THz imaging system in XLPE for power cable insulation[J]. Applied Sciences, 2020, 10(6): 2055.',
        '[16] Li J, et al. Terahertz nondestructive testing method of oil-paper insulation debonding and foreign matter defects[J]. IEEE Transactions on Dielectrics and Electrical Insulation, 2022.',
        '[17] Mei H, Jiang H, Yin F, et al. Detection of small defects in composite insulators using terahertz technique and deconvolution method[J]. IEEE Transactions on Instrumentation and Measurement, 2020, 69(10): 8146-8155.',
        '[18] Xiong W, Li L, Ren J, et al. Terahertz multiple echoes correction and non-destructive testing based on improved wavelet multi-scale analysis[J]. Sensors, 2022, 22(9): 3477.',
        '[19] Yuan H, et al. Moisture transfer characteristics in silicone oil-silicone rubber insulation system considering swelling effect[J]. IEEE Transactions on Dielectrics and Electrical Insulation, 2023, 30(5): 2025-2034.',
        '[20] Huang Z, et al. Detection of moisture content of silicone oil in high voltage cable terminal based on terahertz time-domain spectroscopy[C]//2024 Power System and Green Energy Conference. 2024.',
        '[21] Ellison W J. Permittivity of pure water over the frequency range 0-25 THz and temperature range 0-100 °C[J]. Journal of Physical and Chemical Reference Data, 2007, 36(1): 1-18.',
        '[22] Naito K, Kagawa Y, Utsuno S, et al. Dielectric properties of woven fabric glass fiber reinforced polymer-matrix composites in the THz frequency range[J]. Composites Science and Technology, 2009, 69: 2027-2029.',
        '[23] Withayachumnankul W, Naftaly M. Fundamentals of measurement in terahertz time-domain spectroscopy[J]. Journal of Infrared, Millimeter, and Terahertz Waves, 2014, 35: 610-637.',
        '[24] Dorney T D, Baraniuk R G, Mittleman D M. Material parameter estimation with terahertz time-domain spectroscopy[J]. JOSA A, 2001, 18(7): 1562-1571.',
        '[25] Jepsen P U, Cooke D G, Koch M. Terahertz spectroscopy and imaging—modern techniques and applications[J]. Laser & Photonics Reviews, 2011, 5(1): 124-166.',
        '[26] 周孜毅. 高压电缆瓷套式终端内部绝缘油液面检测技术研究[D]. 华南理工大学, 2019.',
        '[27] 刘东阳. 基于超声波的变压器油位非接触式监测技术研究[D]. 西安科技大学, 2023.',
        '[28] GB/T 11017.3—2014 额定电压110 kV交联聚乙烯绝缘电力电缆及其附件 第3部分：电缆附件[S].',
        '[29] GB/T 18890.3—2015 额定电压220 kV交联聚乙烯绝缘电力电缆及其附件 第3部分：电缆附件[S].',
        '[30] 西安赛谱自动化仪表技术有限公司. RD1200系列线性调频连续波雷达物位计使用说明书[Z].',
        '[31] 项目组. 基于侧壁亚太赫兹回波的高压电缆充油终端油位识别仿真技术报告[R]. 2026.',
        '[32] 项目组. 亚太赫兹探头RS485采集、模拟终端与STM32离线判定实物记录[R]. 2026.',
    ]
    for ref in refs:
        p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(0); p.paragraph_format.first_line_indent=Cm(-0.6); p.paragraph_format.left_indent=Cm(0.6); p.paragraph_format.line_spacing=1.2; p.paragraph_format.space_after=Pt(3)
        r=p.add_run(ref); set_run_font(r,size=9.5)

    # appendix
    doc.add_heading('附录A 前期工作与本开题研究内容对应关系',level=1)
    add_table(doc,['前期工作','已获得认识','在本开题中的处理','后续补强'],[
        ['二维/局部三维COMSOL仿真','空气路径与硅油路径存在幅值和能量差异；相对比值具有稳定性','作为机理可行性和信号处理流程依据，不直接作为目标频点定量结果','重建目标工作频点局部模型、细化网格与边界、与实测回波闭环'],
        ['RS485与220点幅值/CFAR采集','现有探头接口可稳定获取波形与距离字段','算法严格围绕现有可观测量设计，不假设原始I/Q与相位','标定点序号—距离、自动增益影响、波束足迹和质量门限'],
        ['玻纤—硅橡胶模拟终端有油/无油试验','有限固定条件下可见显著波形差异','证明硬件链路可运行，不能等同于界面定位和工程鲁棒性','搭建轴向平台、独立真值、跨日/跨安装/跨容器数据集'],
        ['STM32红绿灯判定','轻量判据能够嵌入运行','作为工程实现基础，不作为算法创新主体','移植经独立测试验证的低维特征、增加置信度与异常状态'],
    ],font_size=8.8)

    filename=OUT/'基于亚太赫兹侧壁回波的电缆充油终端分层识别与界面定位_硕士开题报告_插图完整版.docx'
    doc.save(filename)
    print(filename)


if __name__=='__main__':
    build_document()
