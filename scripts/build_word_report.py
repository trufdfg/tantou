from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIR = Path("outputs")
REPORT_PATH = OUTPUT_DIR / "太赫兹探头油液识别训练报告.docx"


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def pct(value: str | float | int | None, digits: int = 1) -> str:
    if value is None or value == "":
        return "-"
    return f"{float(value) * 100:.{digits}f}%"


def num(value: str | float | int | None, digits: int = 3) -> str:
    if value is None or value == "":
        return "-"
    return f"{float(value):.{digits}f}"


def rel(path: str | Path) -> str:
    path = Path(path)
    try:
        return str(path.resolve().relative_to(Path.cwd().resolve()))
    except ValueError:
        return str(path)


def display_model_name(name: str) -> str:
    return {
        "hist_gradient_boosting": "hist_gbdt",
        "random_forest": "random_forest",
        "extra_trees": "extra_trees",
        "logistic_l2": "logistic_l2",
    }.get(name, name)


def compact_dis_values(values: str) -> str:
    if not values:
        return "-"
    parts = []
    for item in values.split(","):
        item = item.strip()
        if not item:
            continue
        parts.append(str(int(round(float(item)))))
    return ",".join(parts)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(document: Document, headers: list[str], rows: list[list[str]], *, widths: list[float] | None = None) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(header_cells[i], header, bold=True, color="1F2430")
        set_cell_shading(header_cells[i], "F2F4F7")
        if widths:
            header_cells[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if widths:
                cells[i].width = Inches(widths[i])
    document.add_paragraph()


def set_run_font(run, *, size: int | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(level=level)
    run = paragraph.add_run(text)
    if level == 1:
        set_run_font(run, size=16, bold=True, color="2E74B5")
    elif level == 2:
        set_run_font(run, size=13, bold=True, color="2E74B5")
    else:
        set_run_font(run, size=12, bold=True, color="1F4D78")


def add_body_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.1
    run = paragraph.add_run(text)
    set_run_font(run, size=11)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(item)
        set_run_font(run, size=10)


def add_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    set_run_font(run, size=9, color="6F768A")
    run.italic = True


def add_picture(document: Document, path: Path, caption: str) -> None:
    if path.exists():
        document.add_picture(str(path), width=Inches(6.3))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_caption(document, caption)


def style_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def main() -> None:
    summary = json.loads((OUTPUT_DIR / "training_summary.json").read_text(encoding="utf-8"))
    model_scores = read_csv(OUTPUT_DIR / "model_cv_scores.csv")
    distance_scores = read_csv(OUTPUT_DIR / "leave_distance_out_scores.csv")
    test_predictions = read_csv(OUTPUT_DIR / "test_file_predictions.csv")
    feature_importance = read_csv(OUTPUT_DIR / "feature_importance.csv")

    selected = summary["selected_cv"]
    confusion = summary["selected_oof_confusion_matrix"]
    charts = {key: Path(value) for key, value in summary["chart_paths"].items()}

    document = Document()
    style_document(document)

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("太赫兹探头油液识别训练报告")
    set_run_font(run, size=22, bold=True, color="0B2545")

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("基于训练数据的稳健特征提取、模型比较、距离鲁棒性验证与测试集预测")
    set_run_font(run, size=11, color="555555")

    add_heading(document, "技术摘要", 1)
    add_bullets(
        document,
        [
            f"训练数据包含 {summary['train_files']} 个采集文件、{summary['train_blocks']} 个记录块；测试数据包含 {summary['test_files']} 个文件、{summary['test_blocks']} 个记录块。每个记录块包含 Pos、Dis、data 与 CFAR 阈值曲线。",
            f"最终选择 {summary['selected_model_name']}，使用 {summary['selected_feature_count']} 个稳定工程特征，部署阈值为 {summary['selected_threshold']:.2f}。最终特征集中没有直接使用原始 Pos/Dis 作为判别特征，避免模型只学到采集位置差异。",
            f"文件分组交叉验证的文件级 AUC 为 {num(selected['file_auc'], 3)}，文件级平衡准确率为 {pct(selected['file_balanced_accuracy_opt'])}；48 个训练文件的折外预测中，1 个无油文件误判为有油，0 个有油文件漏判。",
            f"按距离留一验证的文件级平衡准确率均值为 {pct(summary['distance_cv_mean_file_balanced_accuracy'])}，最低为 {pct(summary['distance_cv_min_file_balanced_accuracy'])}，说明模型不是单纯依赖某个固定距离或某个峰值。",
            "测试集未提供真实标签，因此不能计算测试准确率；本报告仅给出严格隔离后的测试预测。ceshi_2 的概率为 56.4%，接近阈值，建议后续作为重点复核样本。",
        ],
    )

    add_heading(document, "测试集预测结论", 1)
    test_rows = [
        [
            row["file_id"],
            row["pred_name"],
            pct(row["prob_oil_mean"], 1),
            pct(row["positive_block_frac"], 1),
            row["block_count"],
            compact_dis_values(row["dis_mm_values"]),
        ]
        for row in test_predictions
    ]
    add_table(
        document,
        ["测试文件", "预测", "有油概率", "阳性块比例", "记录块数", "Dis(mm)"],
        test_rows,
        widths=[1.0, 0.6, 0.8, 0.9, 0.7, 2.5],
    )
    add_picture(document, charts["test_predictions"], "图 1. 测试文件有油概率；虚线为部署阈值。")

    add_heading(document, "数据隔离与样本组织", 1)
    add_body_paragraph(
        document,
        "训练数据与测试数据位于不同目录，脚本先解析训练目录完成模型选择、交叉验证、阈值选择和最终拟合；测试目录只在这些步骤完成后读取并预测。"
        "验证时按源文件分组，同一采集文件内的多个记录块不会同时进入训练折和验证折，避免把同一次采集的重复结构泄漏到验证结果中。"
    )
    add_bullets(
        document,
        [
            "训练标签来源：训练数据/有油 与 训练数据/无油 两个目录。",
            "距离层级：0cm、10cm、20cm、30cm、40cm、50cm；每个类别每个距离 4 个训练文件。",
            "基本解析单位：一个 DAT 内的每个 Pos/Dis/data-cfar 块作为记录块，文件级预测通过记录块概率均值聚合。",
            "测试集标签状态：测试文件夹未包含真实标签或标签目录，因此未参与任何训练、特征选择、阈值选择或评价。"
        ],
    )
    add_picture(document, charts["example_waveforms"], "图 2. 有油/无油示例记录块的 data 与 CFAR 曲线。")

    add_heading(document, "特征工程", 1)
    category_rows = [
        ["稳健幅值与形态", str(summary["feature_categories"]["robust_amplitude_shape"]), "中位数/MAD 归一化、分位数、RMS、峰度、偏度、总变差、形状熵"],
        ["CFAR 阈值关系", str(summary["feature_categories"]["cfar_threshold_relation"]), "data-cfar 残差、超过阈值比例、超过阈值面积、最大阈值裕量、data/CFAR 比值"],
        ["多峰结构", str(summary["feature_categories"]["multi_peak_structure"]), "峰数量、峰显著性、峰宽、Top1-Top5 峰位置/幅值/间距、主次峰比例"],
        ["多路径窗口", str(summary["feature_categories"]["window_multipath"]), "将 220 点记录分为 5 个相对距离窗口，提取窗口能量、残差、峰数和前后窗比例"],
        ["频谱结构", str(summary["feature_categories"]["spectral"]), "低/中/高频能量占比、谱质心、谱带宽、谱熵"],
        ["距离归一化", str(summary["feature_categories"]["distance_normalized"]), "相对峰位置、稳健 z-score、能量占比，以及按距离平方校正的幅值/CFAR 裕量"],
    ]
    add_table(document, ["特征组", "数量", "说明"], category_rows, widths=[1.4, 0.6, 4.5])
    add_body_paragraph(
        document,
        "这些特征覆盖界面反射变化、回波幅值变化、多峰结构变化、CFAR 阈值与目标峰关系、探头到容器/液面/背面等多路径回波变化，以及距离变化下的归一化表达。"
        "最终模型没有直接使用原始位置字段，但保留了物理上更合理的相对位置和距离校正幅值。"
    )

    top_features = [[row["feature"], num(row["importance"], 3)] for row in feature_importance[:10]]
    add_table(document, ["Top 特征", "模型重要性"], top_features, widths=[4.5, 1.2])
    add_picture(document, charts["feature_importance"], "图 3. 最终模型中权重最高的工程特征。")

    add_heading(document, "模型与验证方法", 1)
    add_body_paragraph(
        document,
        "候选模型覆盖线性可解释模型和非线性树模型：L2 逻辑回归、ExtraTrees、随机森林、直方梯度提升。"
        "所有模型都在训练集内完成比较；缺失值通过中位数填充，逻辑回归使用标准化。阈值使用文件级折外预测选择，最终阈值为 0.42。"
    )
    stable_rows = []
    for row in model_scores:
        if row["feature_set"] != "stable_no_raw_metadata":
            continue
        stable_rows.append(
            [
                display_model_name(row["model"]),
                pct(row["file_balanced_accuracy_opt"]),
                num(row["file_auc"], 3),
                pct(row["block_balanced_accuracy_opt"]),
                num(row["block_auc"], 3),
            ]
        )
    add_table(
        document,
        ["模型", "文件级平衡准确率", "文件级 AUC", "块级平衡准确率", "块级 AUC"],
        stable_rows,
        widths=[1.6, 1.3, 1.0, 1.3, 1.0],
    )
    add_picture(document, charts["model_comparison"], "图 4. 稳定特征集下的文件分组交叉验证结果。")

    add_heading(document, "距离鲁棒性", 1)
    add_body_paragraph(
        document,
        "为了检验距离归一化是否有效，额外执行按距离留一验证：每次保留一个 nominal distance 作为验证距离，其余五个距离训练模型。"
        "该验证比随机块级划分更接近后续距离变化场景。"
    )
    distance_rows = [
        [
            row["held_out_distance_cm"],
            pct(row["file_balanced_accuracy_at_threshold"]),
            num(row["file_auc"], 3),
            pct(row["block_accuracy_0p5"]),
            row["file_count"],
        ]
        for row in distance_scores
    ]
    add_table(
        document,
        ["留出距离(cm)", "文件级平衡准确率", "文件级 AUC", "块级0.5准确率", "文件数"],
        distance_rows,
        widths=[1.0, 1.5, 1.0, 1.4, 0.8],
    )
    add_picture(document, charts["distance_validation"], "图 5. 按距离留一验证；40cm 与 50cm 相对更弱，但仍高于 87.5%。")

    add_heading(document, "折外误差与阈值", 1)
    add_table(
        document,
        ["真实/预测", "预测无油", "预测有油"],
        [
            ["真实无油", str(confusion["true_no_oil_pred_no_oil"]), str(confusion["true_no_oil_pred_oil"])],
            ["真实有油", str(confusion["true_oil_pred_no_oil"]), str(confusion["true_oil_pred_oil"])],
        ],
        widths=[1.4, 1.2, 1.2],
    )
    add_body_paragraph(
        document,
        f"折外文件级混淆矩阵显示，当前阈值下有油样本无漏判；唯一错误是一个无油训练文件被判为有油。"
        f"对于安全或防漏油场景，可以接受略偏保守的阈值；如果后续更关注误报率，可以在真实测试标签回填后重新调整阈值。"
    )

    add_heading(document, "局限性与后续建议", 1)
    add_bullets(
        document,
        [
            "当前测试集没有标签，无法报告真实测试准确率。拿到测试标签后，应只用 test_file_predictions.csv 做一次性评价，不要再用测试结果反向调模型。",
            "训练集中每个距离每类只有 4 个文件，文件数仍偏少；虽然块数较多，但同一文件内块高度相关，因此报告优先使用文件级指标。",
            "ceshi_2 的预测概率接近阈值且块内一致性不高，建议复测或人工复核原始波形。",
            "后续应加入真实容器、不同容器材料/壁厚、不同油量高度、重复装夹、温漂与探头姿态变化数据，继续做文件级或批次级分组验证。",
            "若将模型部署到新采集批次，建议先用少量已知标签样本校准阈值，并监控 CFAR 残差、窗口能量和多峰间距的分布漂移。",
        ],
    )

    add_heading(document, "可复现产物", 1)
    add_table(
        document,
        ["产物", "路径"],
        [
            ["训练脚本", rel("scripts/train_oil_classifier.py")],
            ["报告脚本", rel("scripts/build_word_report.py")],
            ["最终模型", rel(summary["model_path"])],
            ["测试文件级预测", rel(OUTPUT_DIR / "test_file_predictions.csv")],
            ["训练摘要", rel(OUTPUT_DIR / "training_summary.json")],
            ["本报告", rel(REPORT_PATH)],
        ],
        widths=[1.5, 5.0],
    )
    add_body_paragraph(
        document,
        "复现命令：python scripts/train_oil_classifier.py --data-root shuju --output-dir outputs；"
        "随后运行 python scripts/build_word_report.py 生成本 Word 报告。"
    )

    document.save(REPORT_PATH)
    print(REPORT_PATH.resolve())


if __name__ == "__main__":
    main()
